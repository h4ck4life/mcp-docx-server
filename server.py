from mcp.server.fastmcp import FastMCP, Context
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR
from io import BytesIO
import base64
import os
from docx2pdf import convert
from docx.enum.section import WD_SECTION, WD_ORIENT

# Create an MCP server specifically for Word document operations
mcp = FastMCP("WordDocServer", 
              description="An MCP server that allows reading and manipulating Microsoft Word (.docx) files. "
                          "This server can create, read, and modify Word documents stored in the same directory as the script.")

# Helper functions
def get_document_path(doc_id: str) -> str:
    """Returns the full path to a document in the same directory as this script."""
    script_dir = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(script_dir, f"{doc_id}.docx")

def load_document(doc_id: str) -> Document:
    """Loads a Word document, handling potential FileNotFoundError."""
    doc_path = get_document_path(doc_id)
    try:
        return Document(doc_path)
    except FileNotFoundError:
        raise ValueError(f"Document '{doc_id}.docx' not found.")
    except Exception as e:
        raise ValueError(f"Error loading document '{doc_id}.docx': {str(e)}")

# Check if a style exists in a document
def style_exists(document, style_name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    """Checks if a style exists in the document."""
    for style in document.styles:
        if style.name == style_name and style.type == style_type:
            return True
    return False

# Document reading operations
@mcp.resource("word://{doc_id}/content")
def get_document_content(doc_id: str) -> str:
    """Reads the content of a Microsoft Word (.docx) document and returns it as text."""
    try:
        document = load_document(doc_id)
        full_text = [paragraph.text for paragraph in document.paragraphs]
        return '\n'.join(full_text)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Unexpected error: {str(e)}"

@mcp.tool()
def read_document(doc_id: str) -> str:
    """Reads the entire content of a Word document."""
    return get_document_content(doc_id)

@mcp.tool()
def check_document_exists(doc_id: str) -> str:
    """Checks if a Word document exists and can be read."""
    doc_path = get_document_path(doc_id)
    try:
        if os.path.exists(doc_path):
            document = Document(doc_path)
            paragraph_count = len(document.paragraphs)
            return f"Document '{doc_id}.docx' exists and is readable at path: {os.path.abspath(doc_path)}. Contains {paragraph_count} paragraphs."
        else:
            return f"Document '{doc_id}.docx' does not exist at path: {os.path.abspath(doc_path)}"
    except Exception as e:
        return f"Document '{doc_id}.docx' exists but cannot be read: {str(e)}"

@mcp.tool()
def list_available_documents() -> str:
    """Lists all Word documents (.docx files) available in the server directory."""
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        docx_files = [f.replace('.docx', '') for f in os.listdir(script_dir) if f.endswith('.docx')]
        
        if not docx_files:
            return "No Word documents (.docx files) found in the server directory."
        
        doc_list = "\n".join([f"- {doc}" for doc in docx_files])
        return f"Available Word documents (without .docx extension):\n{doc_list}"
    except Exception as e:
        return f"Error listing documents: {str(e)}"

# New Style Management Functions
@mcp.tool()
def ensure_style_exists(doc_id: str, style_name: str, style_type: str = "paragraph") -> str:
    """
    Ensures a style exists in the document by applying it to a temporary paragraph.
    
    This is useful for built-in styles that need to be defined in the document before use.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        style_name (str): The name of the style to ensure exists.
        style_type (str): Type of style: 'paragraph', 'character', or 'table'.
    
    Returns:
        str: Status message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        # Map string type to enum
        style_type_map = {
            "paragraph": WD_STYLE_TYPE.PARAGRAPH,
            "character": WD_STYLE_TYPE.CHARACTER,
            "table": WD_STYLE_TYPE.TABLE
        }
        
        if style_type.lower() not in style_type_map:
            return f"Error: Invalid style type '{style_type}'. Valid values are: {', '.join(style_type_map.keys())}"
        
        style_type_enum = style_type_map[style_type.lower()]
        
        # Check if style already exists
        if style_exists(document, style_name, style_type_enum):
            return f"Style '{style_name}' already exists in document."
        
        # For paragraph style, add a temporary paragraph with the style
        if style_type_enum == WD_STYLE_TYPE.PARAGRAPH:
            try:
                temp_para = document.add_paragraph("Style Definition", style=style_name)
                # Remove the temporary paragraph
                p = temp_para._element
                p.getparent().remove(p)
                
                doc_path = get_document_path(doc_id)
                document.save(doc_path)
                return f"Paragraph style '{style_name}' successfully defined in document."
            except KeyError:
                return f"Error: Built-in style '{style_name}' not found in Word. Check the style name."
        
        # For character style, add a paragraph with a run using the style
        elif style_type_enum == WD_STYLE_TYPE.CHARACTER:
            try:
                temp_para = document.add_paragraph()
                temp_run = temp_para.add_run("Style Definition")
                temp_run.style = style_name
                
                # Remove the temporary paragraph
                p = temp_para._element
                p.getparent().remove(p)
                
                doc_path = get_document_path(doc_id)
                document.save(doc_path)
                return f"Character style '{style_name}' successfully defined in document."
            except KeyError:
                return f"Error: Built-in style '{style_name}' not found in Word. Check the style name."
        
        # For table style, add a temporary table with the style
        elif style_type_enum == WD_STYLE_TYPE.TABLE:
            try:
                temp_table = document.add_table(rows=1, cols=1)
                temp_table.style = style_name
                
                # Remove the temporary table
                p = temp_table._element
                p.getparent().remove(p)
                
                doc_path = get_document_path(doc_id)
                document.save(doc_path)
                return f"Table style '{style_name}' successfully defined in document."
            except KeyError:
                return f"Error: Built-in style '{style_name}' not found in Word. Check the style name."
        
        return "Unknown error ensuring style exists."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error ensuring style exists: {str(e)}"

@mcp.tool()
def create_custom_style(doc_id: str, style_name: str, style_type: str = "paragraph", base_style: str = None) -> str:
    """
    Creates a custom style in the document.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        style_name (str): The name for the new style.
        style_type (str): Type of style: 'paragraph', 'character', or 'table'.
        base_style (str, optional): Name of the style to base this one on.
    
    Returns:
        str: Status message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        # Map string type to enum
        style_type_map = {
            "paragraph": WD_STYLE_TYPE.PARAGRAPH,
            "character": WD_STYLE_TYPE.CHARACTER,
            "table": WD_STYLE_TYPE.TABLE
        }
        
        if style_type.lower() not in style_type_map:
            return f"Error: Invalid style type '{style_type}'. Valid values are: {', '.join(style_type_map.keys())}"
        
        style_type_enum = style_type_map[style_type.lower()]
        
        # Check if style already exists
        if style_exists(document, style_name, style_type_enum):
            return f"Error: Style '{style_name}' already exists in document."
        
        # If base_style is provided, ensure it exists first
        if base_style:
            if not style_exists(document, base_style, style_type_enum):
                # Try to ensure base style exists if it's a built-in style
                ensure_result = ensure_style_exists(doc_id, base_style, style_type)
                if "Error" in ensure_result or "not found" in ensure_result:
                    return f"Error: Base style '{base_style}' does not exist and could not be defined."
        
        # Create the new style
        new_style = document.styles.add_style(style_name, style_type_enum)
        
        # Set base style if provided
        if base_style:
            try:
                new_style.base_style = document.styles[base_style]
            except KeyError:
                return f"Error setting base style: Style '{base_style}' not found."
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Custom {style_type} style '{style_name}' created successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error creating custom style: {str(e)}"

@mcp.tool()
def modify_style(doc_id: str, style_name: str, properties: dict) -> str:
    """
    Modifies properties of an existing style.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        style_name (str): The name of the style to modify.
        properties (dict): Dictionary with style properties to modify:
            - font: Font properties dict (name, size, bold, italic, underline, color)
            - paragraph: Paragraph formatting properties dict
    
    Returns:
        str: Status message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        # Check if style exists
        try:
            style = document.styles[style_name]
        except KeyError:
            return f"Error: Style '{style_name}' not found in document."
        
        # Modify font properties if provided
        font_props = properties.get("font", {})
        if font_props and hasattr(style, "font"):
            font = style.font
            
            # Font name
            if "name" in font_props:
                font.name = font_props["name"]
            
            # Font size
            if "size" in font_props:
                font.size = Pt(float(font_props["size"]))
            
            # Font styles
            if "bold" in font_props:
                font.bold = bool(font_props["bold"])
            
            if "italic" in font_props:
                font.italic = bool(font_props["italic"])
            
            if "underline" in font_props:
                font.underline = bool(font_props["underline"])
            
            # Font color
            if "color" in font_props:
                color = font_props["color"]
                if color.startswith('#'):
                    # Convert hex color to RGB
                    r = int(color[1:3], 16)
                    g = int(color[3:5], 16)
                    b = int(color[5:7], 16)
                    font.color.rgb = RGBColor(r, g, b)
                elif color.startswith('rgb('):
                    # Parse rgb() format
                    rgb = color.strip('rgb()').split(',')
                    if len(rgb) == 3:
                        r = int(rgb[0].strip())
                        g = int(rgb[1].strip())
                        b = int(rgb[2].strip())
                        font.color.rgb = RGBColor(r, g, b)
        
        # Modify paragraph formatting properties if provided
        para_props = properties.get("paragraph", {})
        if para_props and hasattr(style, "paragraph_format"):
            para_format = style.paragraph_format
            
            # Alignment
            alignment = para_props.get("alignment")
            if alignment:
                alignment_map = {
                    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if alignment.upper() in alignment_map:
                    para_format.alignment = alignment_map[alignment.upper()]
            
            # Indentation
            left_indent = para_props.get("left_indent")
            if left_indent is not None:
                para_format.left_indent = Inches(float(left_indent))
            
            right_indent = para_props.get("right_indent")
            if right_indent is not None:
                para_format.right_indent = Inches(float(right_indent))
            
            first_line_indent = para_props.get("first_line_indent")
            if first_line_indent is not None:
                para_format.first_line_indent = Inches(float(first_line_indent))
            
            # Spacing
            space_before = para_props.get("space_before")
            if space_before is not None:
                para_format.space_before = Pt(float(space_before))
            
            space_after = para_props.get("space_after")
            if space_after is not None:
                para_format.space_after = Pt(float(space_after))
            
            # Line spacing
            line_spacing = para_props.get("line_spacing")
            if line_spacing is not None:
                try:
                    # Try as a float (multiple)
                    spacing_float = float(line_spacing)
                    para_format.line_spacing = spacing_float
                except ValueError:
                    # Try as a point value
                    para_format.line_spacing = Pt(float(line_spacing))
            
            # Pagination
            keep_together = para_props.get("keep_together")
            if keep_together is not None:
                para_format.keep_together = bool(keep_together)
            
            keep_with_next = para_props.get("keep_with_next")
            if keep_with_next is not None:
                para_format.keep_with_next = bool(keep_with_next)
            
            page_break_before = para_props.get("page_break_before")
            if page_break_before is not None:
                para_format.page_break_before = bool(page_break_before)
            
            widow_control = para_props.get("widow_control")
            if widow_control is not None:
                para_format.widow_control = bool(widow_control)
        
        # Additional style properties
        if "quick_style" in properties:
            style.quick_style = bool(properties["quick_style"])
        
        if "hidden" in properties:
            style.hidden = bool(properties["hidden"])
        
        if "priority" in properties:
            style.priority = int(properties["priority"])
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Style '{style_name}' modified successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error modifying style: {str(e)}"

@mcp.tool()
def get_styles_detail(doc_id: str, style_type: str = None) -> str:
    """
    Gets detailed information about styles in the document.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        style_type (str, optional): Filter by style type: 'paragraph', 'character', or 'table'.
                                   If None, shows all styles.
    
    Returns:
        str: Detailed information about styles in the document.
    """
    try:
        document = load_document(doc_id)
        
        # Map string type to enum if provided
        style_type_enum = None
        if style_type:
            style_type_map = {
                "paragraph": WD_STYLE_TYPE.PARAGRAPH,
                "character": WD_STYLE_TYPE.CHARACTER,
                "table": WD_STYLE_TYPE.TABLE
            }
            
            if style_type.lower() not in style_type_map:
                return f"Error: Invalid style type '{style_type}'. Valid values are: {', '.join(style_type_map.keys())}"
            
            style_type_enum = style_type_map[style_type.lower()]
        
        # Get styles
        styles_info = []
        for style in document.styles:
            # Skip if filtering by type and this style doesn't match
            if style_type_enum and style.type != style_type_enum:
                continue
            
            # Get style type as string
            style_type_str = "Unknown"
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                style_type_str = "Paragraph"
            elif style.type == WD_STYLE_TYPE.CHARACTER:
                style_type_str = "Character"
            elif style.type == WD_STYLE_TYPE.TABLE:
                style_type_str = "Table"
            elif style.type == WD_STYLE_TYPE.LIST:
                style_type_str = "List"
            
            # Get base style name if available
            base_style = "None"
            if hasattr(style, "base_style") and style.base_style:
                base_style = style.base_style.name
            
            # Create style info
            style_info = [
                f"Style: {style.name}",
                f"  Type: {style_type_str}",
                f"  Base Style: {base_style}"
            ]
            
            # Add behavior properties if available
            behavior_props = []
            if hasattr(style, "quick_style"):
                behavior_props.append(f"Quick Style: {style.quick_style}")
            if hasattr(style, "priority"):
                behavior_props.append(f"Priority: {style.priority}")
            if hasattr(style, "hidden"):
                behavior_props.append(f"Hidden: {style.hidden}")
            
            if behavior_props:
                style_info.append("  Behavior:")
                for prop in behavior_props:
                    style_info.append(f"    {prop}")
            
            styles_info.append("\n".join(style_info))
        
        if not styles_info:
            return f"No styles found in document{' with type ' + style_type if style_type else ''}."
        
        return "\n\n".join(styles_info)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error getting styles detail: {str(e)}"

@mcp.tool()
def check_style_usage(doc_id: str, style_name: str) -> str:
    """
    Checks if a style is used in the document and where.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        style_name (str): The name of the style to check.
    
    Returns:
        str: Information about where the style is used.
    """
    try:
        document = load_document(doc_id)
        
        # Check if style exists
        try:
            style = document.styles[style_name]
        except KeyError:
            return f"Style '{style_name}' not found in document."
        
        # Get style type
        style_type_str = "Unknown"
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style_type_str = "Paragraph"
        elif style.type == WD_STYLE_TYPE.CHARACTER:
            style_type_str = "Character"
        elif style.type == WD_STYLE_TYPE.TABLE:
            style_type_str = "Table"
        
        # Check usage based on style type
        usage_locations = []
        
        # Check paragraphs for paragraph and character styles
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
            for i, para in enumerate(document.paragraphs):
                # Check paragraph style
                if style.type == WD_STYLE_TYPE.PARAGRAPH and para.style and para.style.name == style_name:
                    preview = para.text[:30] + ("..." if len(para.text) > 30 else "")
                    usage_locations.append(f"Paragraph {i}: \"{preview}\"")
                
                # Check character styles in runs
                if style.type == WD_STYLE_TYPE.CHARACTER:
                    for j, run in enumerate(para.runs):
                        if run.style and run.style.name == style_name:
                            preview = run.text[:30] + ("..." if len(run.text) > 30 else "")
                            usage_locations.append(f"Paragraph {i}, Run {j}: \"{preview}\"")
        
        # Check tables for table styles
        if style.type == WD_STYLE_TYPE.TABLE:
            for i, table in enumerate(document.tables):
                if table.style and table.style.name == style_name:
                    rows = len(table.rows)
                    cols = len(table.rows[0].cells) if rows > 0 else 0
                    usage_locations.append(f"Table {i}: {rows}x{cols} table")
        
        # Report results
        if not usage_locations:
            return f"{style_type_str} style '{style_name}' exists in the document but is not currently used."
        
        return f"{style_type_str} style '{style_name}' is used in the following locations:\n" + "\n".join(usage_locations)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error checking style usage: {str(e)}"

# Document creation and update operations
def _add_content_to_document(document, content):
    """Helper function to add content to a document object."""
    if not content:
        return True
        
    for item in content:
        content_type = item.get("type", "").lower()
        text = item.get("text", "")
        
        if content_type == "heading":
            level = item.get("level", 1)
            heading = document.add_heading(text, level)
            
            # Apply formatting if provided
            formatting = item.get("formatting", {})
            if formatting:
                _apply_paragraph_formatting(heading, formatting)
            
        elif content_type == "paragraph":
            style = item.get("style")
            try:
                # If style is specified, ensure it exists in the document
                if style:
                    style_exists_in_doc = False
                    for doc_style in document.styles:
                        if doc_style.name == style:
                            style_exists_in_doc = True
                            break
                    
                    # If style doesn't exist, it might be a built-in style that needs defining
                    if not style_exists_in_doc:
                        # Add temporary paragraph to define the style
                        temp_para = document.add_paragraph("", style=style)
                        # Remove the temporary paragraph
                        p = temp_para._element
                        p.getparent().remove(p)
                
                paragraph = document.add_paragraph(text, style=style if style else None)
                
                # Apply formatting if provided
                formatting = item.get("formatting", {})
                if formatting:
                    _apply_paragraph_formatting(paragraph, formatting)
                
                # Apply run formatting if provided
                run_formatting = item.get("run_formatting", {})
                if run_formatting and len(paragraph.runs) > 0:
                    _apply_run_formatting(paragraph.runs[0], run_formatting)
                
            except KeyError:
                # Style not found, add without style
                paragraph = document.add_paragraph(text)
                
        elif content_type == "table":
            rows = item.get("rows", 1)
            cols = item.get("cols", 1)
            data = item.get("data", "")
            style = item.get("style")
            
            table = document.add_table(rows=rows, cols=cols)
            
            # Apply style if specified
            if style:
                try:
                    # If style doesn't exist, it might be a built-in style that needs defining
                    style_exists_in_doc = False
                    for doc_style in document.styles:
                        if doc_style.name == style and doc_style.type == WD_STYLE_TYPE.TABLE:
                            style_exists_in_doc = True
                            break
                    
                    # If style doesn't exist, it might be a built-in style that needs defining
                    if not style_exists_in_doc:
                        # Add temporary table to define the style
                        temp_table = document.add_table(rows=1, cols=1)
                        temp_table.style = style
                        # Remove the temporary table
                        p = temp_table._element
                        p.getparent().remove(p)
                    
                    table.style = style
                except KeyError:
                    pass  # Continue without style if not found
            
            # Fill with data if provided
            if data:
                data_list = data.split(',')
                
                # Pad with empty strings if too few data elements
                if len(data_list) < rows * cols:
                    data_list.extend([''] * (rows * cols - len(data_list)))
                    
                # Check if data matches table dimensions
                if len(data_list) != rows * cols:
                    return False
                
                # Fill table cells
                for i in range(rows):
                    for j in range(cols):
                        cell_idx = i * cols + j
                        if cell_idx < len(data_list):
                            table.cell(i, j).text = data_list[cell_idx].strip()
            
            # Process cell_formatting if provided
            cell_formatting = item.get("cell_formatting", [])
            for cell_format in cell_formatting:
                row = cell_format.get("row", 0)
                col = cell_format.get("col", 0)
                formatting = cell_format.get("formatting", {})
                
                if row < rows and col < cols:
                    cell = table.cell(row, col)
                    if cell and len(cell.paragraphs) > 0:
                        _apply_paragraph_formatting(cell.paragraphs[0], formatting)
    
    return True

def _apply_paragraph_formatting(paragraph, formatting):
    """Apply formatting to a paragraph."""
    if not formatting:
        return
    
    para_format = paragraph.paragraph_format
    
    # Alignment
    alignment = formatting.get("alignment")
    if alignment:
        alignment_map = {
            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
            "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if alignment.upper() in alignment_map:
            para_format.alignment = alignment_map[alignment.upper()]
    
    # Indentation
    left_indent = formatting.get("left_indent")
    if left_indent is not None:
        para_format.left_indent = Inches(float(left_indent))
    
    right_indent = formatting.get("right_indent")
    if right_indent is not None:
        para_format.right_indent = Inches(float(right_indent))
    
    first_line_indent = formatting.get("first_line_indent")
    if first_line_indent is not None:
        para_format.first_line_indent = Inches(float(first_line_indent))
    
    # Spacing
    space_before = formatting.get("space_before")
    if space_before is not None:
        para_format.space_before = Pt(float(space_before))
    
    space_after = formatting.get("space_after")
    if space_after is not None:
        para_format.space_after = Pt(float(space_after))
    
    # Line spacing
    line_spacing = formatting.get("line_spacing")
    if line_spacing is not None:
        try:
            # Try as a float (multiple)
            spacing_float = float(line_spacing)
            para_format.line_spacing = spacing_float
        except ValueError:
            # Try as a point value
            para_format.line_spacing = Pt(float(line_spacing))
    
    # Pagination
    keep_together = formatting.get("keep_together")
    if keep_together is not None:
        para_format.keep_together = bool(keep_together)
    
    keep_with_next = formatting.get("keep_with_next")
    if keep_with_next is not None:
        para_format.keep_with_next = bool(keep_with_next)
    
    page_break_before = formatting.get("page_break_before")
    if page_break_before is not None:
        para_format.page_break_before = bool(page_break_before)
    
    widow_control = formatting.get("widow_control")
    if widow_control is not None:
        para_format.widow_control = bool(widow_control)

def _apply_run_formatting(run, formatting):
    """Apply formatting to a run of text."""
    if not formatting:
        return
    
    font = run.font
    
    # Font name and size
    font_name = formatting.get("name")
    if font_name:
        font.name = font_name
    
    font_size = formatting.get("size")
    if font_size is not None:
        font.size = Pt(float(font_size))
    
    # Font styles
    bold = formatting.get("bold")
    if bold is not None:
        font.bold = bool(bold)
    
    italic = formatting.get("italic")
    if italic is not None:
        font.italic = bool(italic)
    
    underline = formatting.get("underline")
    if underline is not None:
        font.underline = bool(underline)
    
    # Font color
    color = formatting.get("color")
    if color:
        if color.startswith('#'):
            # Convert hex color to RGB
            r = int(color[1:3], 16)
            g = int(color[3:5], 16)
            b = int(color[5:7], 16)
            font.color.rgb = RGBColor(r, g, b)
        elif color.startswith('rgb('):
            # Parse rgb() format
            rgb = color.strip('rgb()').split(',')
            if len(rgb) == 3:
                r = int(rgb[0].strip())
                g = int(rgb[1].strip())
                b = int(rgb[2].strip())
                font.color.rgb = RGBColor(r, g, b)

@mcp.tool()
def create_document(doc_id: str, title: str = "New Document") -> str:
    """Creates a new Word document with a title."""
    try:
        document = Document()
        document.add_heading(title, 0)
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Document '{doc_id}.docx' created successfully at path: {os.path.abspath(doc_path)}"
    except Exception as e:
        return f"Error creating document: {str(e)}"

@mcp.tool()
def create_complete_document(doc_id: str, title: str = "New Document", content: list = None) -> str:
    """Creates a new Word document with title and content in a single operation."""
    try:
        document = Document()
        document.add_heading(title, 0)
        
        if not _add_content_to_document(document, content):
            return "Error in table data: Number of data elements does not match table dimensions."
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        
        return f"Document '{doc_id}.docx' created successfully with title and {len(content) if content else 0} content items at path: {os.path.abspath(doc_path)}"
    except Exception as e:
        return f"Error creating document: {str(e)}"

@mcp.tool()
def update_document(doc_id: str, title: str = None, content: list = None, append: bool = True) -> str:
    """Updates an existing Word document by appending or replacing content."""
    try:
        doc_path = get_document_path(doc_id)
        
        if not os.path.exists(doc_path) and append:
            return f"Document '{doc_id}.docx' does not exist and cannot be updated. Create it first."
        
        if not append or not os.path.exists(doc_path):
            document = Document()
            if title:
                document.add_heading(title, 0)
        else:
            document = Document(doc_path)
            if title:
                document.add_heading(title, 1)
        
        if not _add_content_to_document(document, content):
            return "Error in table data: Number of data elements does not match table dimensions."
        
        document.save(doc_path)
        
        action = "updated by appending" if append else "replaced"
        title_msg = f" with new title" if title else ""
        content_msg = f" and {len(content) if content else 0} content items" if content else ""
        
        return f"Document '{doc_id}.docx' {action}{title_msg}{content_msg} successfully."
    except Exception as e:
        return f"Error updating document: {str(e)}"

@mcp.tool()
def append_to_document(doc_id: str, content: list) -> str:
    """Appends content to an existing Word document."""
    return update_document(doc_id, title=None, content=content, append=True)

@mcp.tool()
def replace_document(doc_id: str, title: str = None, content: list = None) -> str:
    """Replaces an existing Word document with new content."""
    return update_document(doc_id, title=title, content=content, append=False)

@mcp.tool()
def add_section(doc_id: str, start_type: str = "NEW_PAGE") -> str:
    """Adds a new section to the end of a document.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        start_type (str): The type of section break, one of:
            - "NEW_PAGE" (default) - Start the section on a new page
            - "EVEN_PAGE" - Start the section on the next even-numbered page
            - "ODD_PAGE" - Start the section on the next odd-numbered page
            - "CONTINUOUS" - No page break, continue on same page
    
    Returns:
        str: A message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        # Map string values to WD_SECTION enum
        section_types = {
            "NEW_PAGE": WD_SECTION.NEW_PAGE,
            "EVEN_PAGE": WD_SECTION.EVEN_PAGE, 
            "ODD_PAGE": WD_SECTION.ODD_PAGE,
            "CONTINUOUS": WD_SECTION.CONTINUOUS
        }
        
        # Get the section type value
        section_type = section_types.get(start_type.upper())
        if not section_type:
            return f"Error: Invalid section start type '{start_type}'. Valid values are: {', '.join(section_types.keys())}"
        
        # Add the new section
        document.add_section(section_type)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Section with start type '{start_type}' added successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding section: {str(e)}"

@mcp.tool()
def list_sections(doc_id: str) -> str:
    """Lists all sections in a document with their properties.
    
    Args:
        doc_id (str): The document ID (filename without extension).
    
    Returns:
        str: Information about each section in the document.
    """
    try:
        document = load_document(doc_id)
        
        if not document.sections:
            return f"No sections found in document '{doc_id}.docx'."
        
        sections_info = []
        for i, section in enumerate(document.sections):
            # Map orientation value to readable string
            orientation = "PORTRAIT" if section.orientation == WD_ORIENT.PORTRAIT else "LANDSCAPE"
            
            # Convert dimensions to inches for readability
            page_width_inches = section.page_width / 914400  # 914400 = 1 inch in EMUs
            page_height_inches = section.page_height / 914400
            
            # Create info string
            section_info = [
                f"Section {i}:",
                f"  Start Type: {section.start_type}",
                f"  Orientation: {orientation}",
                f"  Page Size: {page_width_inches:.2f}\" x {page_height_inches:.2f}\"",
                f"  Margins (inches):",
                f"    Left: {section.left_margin/914400:.2f}\"",
                f"    Right: {section.right_margin/914400:.2f}\"",
                f"    Top: {section.top_margin/914400:.2f}\"",
                f"    Bottom: {section.bottom_margin/914400:.2f}\"",
                f"    Gutter: {section.gutter/914400:.2f}\"",
                f"    Header Distance: {section.header_distance/914400:.2f}\"",
                f"    Footer Distance: {section.footer_distance/914400:.2f}\""
            ]
            sections_info.append("\n".join(section_info))
        
        return "\n\n".join(sections_info)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error listing sections: {str(e)}"

@mcp.tool()
def set_section_properties(doc_id: str, section_index: int, properties: dict) -> str:
    """Sets properties for a specific section in the document.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        section_index (int): The index of the section to modify (0-based).
        properties (dict): Dictionary with section properties:
            - start_type: Section break type ("NEW_PAGE", "EVEN_PAGE", "ODD_PAGE", "CONTINUOUS")
            - orientation: "PORTRAIT" or "LANDSCAPE"
            - page_width: Page width in inches
            - page_height: Page height in inches
            - left_margin, right_margin, top_margin, bottom_margin: Margins in inches
            - gutter: Gutter margin in inches
            - header_distance, footer_distance: Header/footer distance in inches
    
    Returns:
        str: A message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        if not document.sections or section_index >= len(document.sections):
            return f"Error: Section index {section_index} is out of range. Document has {len(document.sections) if document.sections else 0} sections."
        
        section = document.sections[section_index]
        
        # Handle start_type
        if "start_type" in properties:
            start_type = properties["start_type"].upper()
            section_types = {
                "NEW_PAGE": WD_SECTION.NEW_PAGE,
                "EVEN_PAGE": WD_SECTION.EVEN_PAGE, 
                "ODD_PAGE": WD_SECTION.ODD_PAGE,
                "CONTINUOUS": WD_SECTION.CONTINUOUS
            }
            if start_type in section_types:
                section.start_type = section_types[start_type]
            else:
                return f"Error: Invalid section start type '{start_type}'. Valid values are: {', '.join(section_types.keys())}"
        
        # Handle orientation
        if "orientation" in properties:
            orientation = properties["orientation"].upper()
            if orientation == "LANDSCAPE":
                # If changing to landscape, may need to swap width and height
                if section.orientation == WD_ORIENT.PORTRAIT:
                    # Store current dimensions before changing orientation
                    old_width, old_height = section.page_width, section.page_height
                    # Set orientation first
                    section.orientation = WD_ORIENT.LANDSCAPE
                    # If page dimensions are not explicitly set in properties, swap them
                    if "page_width" not in properties and "page_height" not in properties:
                        section.page_width, section.page_height = old_height, old_width
                else:
                    # Already landscape, just ensure orientation is set
                    section.orientation = WD_ORIENT.LANDSCAPE
            elif orientation == "PORTRAIT":
                # If changing to portrait, may need to swap width and height
                if section.orientation == WD_ORIENT.LANDSCAPE:
                    # Store current dimensions before changing orientation
                    old_width, old_height = section.page_width, section.page_height
                    # Set orientation first
                    section.orientation = WD_ORIENT.PORTRAIT
                    # If page dimensions are not explicitly set in properties, swap them
                    if "page_width" not in properties and "page_height" not in properties:
                        section.page_width, section.page_height = old_height, old_width
                else:
                    # Already portrait, just ensure orientation is set
                    section.orientation = WD_ORIENT.PORTRAIT
            else:
                return f"Error: Invalid orientation '{orientation}'. Valid values are: PORTRAIT, LANDSCAPE"
        
        # Handle page dimensions (after orientation changes, if any)
        if "page_width" in properties:
            section.page_width = int(float(properties["page_width"]) * 914400)  # Convert inches to EMUs
        
        if "page_height" in properties:
            section.page_height = int(float(properties["page_height"]) * 914400)  # Convert inches to EMUs
        
        # Handle margins
        for margin_prop in ["left_margin", "right_margin", "top_margin", "bottom_margin", 
                           "gutter", "header_distance", "footer_distance"]:
            if margin_prop in properties:
                setattr(section, margin_prop, int(float(properties[margin_prop]) * 914400))  # Convert inches to EMUs
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Properties for section {section_index} updated successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error setting section properties: {str(e)}"

@mcp.tool()
def change_page_orientation(doc_id: str, section_index: int, orientation: str) -> str:
    """Changes the page orientation for a specific section.
    
    This is a convenience function that wraps set_section_properties for
    the common task of changing page orientation.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        section_index (int): The index of the section to modify (0-based).
        orientation (str): "PORTRAIT" or "LANDSCAPE"
    
    Returns:
        str: A message indicating success or failure.
    """
    return set_section_properties(doc_id, section_index, {"orientation": orientation})

# Individual content addition operations
@mcp.tool()
def add_paragraph(doc_id: str, text: str, style: str = None, formatting: dict = None) -> str:
    """Adds a paragraph to an existing Word document, optionally with style and formatting.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        text (str): The paragraph text.
        style (str, optional): The paragraph style.
        formatting (dict, optional): Dictionary with paragraph formatting options like:
            - alignment: 'LEFT', 'CENTER', 'RIGHT', 'JUSTIFY'
            - left_indent, right_indent, first_line_indent: Indentation in inches
            - space_before, space_after: Spacing in points
            - line_spacing: Line spacing as multiple or points
            - keep_together, keep_with_next, page_break_before, widow_control: Boolean pagination options
    """
    try:
        document = load_document(doc_id)
        result_message = "Paragraph added successfully."
        
        # If style is specified, ensure it exists in the document
        if style:
            style_exists_in_doc = False
            for doc_style in document.styles:
                if doc_style.name == style and doc_style.type == WD_STYLE_TYPE.PARAGRAPH:
                    style_exists_in_doc = True
                    break
            
            # If style doesn't exist, it might be a built-in style that needs defining
            if not style_exists_in_doc:
                try:
                    # Add temporary paragraph to define the style
                    temp_para = document.add_paragraph("", style=style)
                    # Remove the temporary paragraph
                    p = temp_para._element
                    p.getparent().remove(p)
                except KeyError:
                    result_message = f"Warning: Style '{style}' not found. Added without style."
                    style = None
        
        # Now add the actual paragraph
        if style:
            try:
                paragraph = document.add_paragraph(text, style=style)
            except KeyError:
                paragraph = document.add_paragraph(text)
                result_message = f"Warning: Style '{style}' not found. Added without style."
        else:
            paragraph = document.add_paragraph(text)
        
        # Apply formatting if provided
        if formatting:
            _apply_paragraph_formatting(paragraph, formatting)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return result_message
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding paragraph: {str(e)}"

@mcp.tool()
def add_formatted_text(doc_id: str, paragraph_index: int, text: str, formatting: dict = None) -> str:
    """Adds formatted text to an existing paragraph.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        paragraph_index (int): The index of the paragraph to add text to (0-based).
        text (str): The text to add.
        formatting (dict, optional): Dictionary with text formatting options like:
            - name: Font name
            - size: Font size in points
            - bold, italic, underline: Boolean style options
            - color: Color as hex (#RRGGBB) or rgb(r,g,b)
    """
    try:
        document = load_document(doc_id)
        
        if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
            return "Error: Paragraph index out of range."
        
        paragraph = document.paragraphs[paragraph_index]
        run = paragraph.add_run(text)
        
        # Apply formatting if provided
        if formatting:
            _apply_run_formatting(run, formatting)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return "Formatted text added successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding formatted text: {str(e)}"

@mcp.tool()
def add_image(doc_id: str, image_data: str, image_name: str, width_inches: float = 6.0) -> str:
    """Adds an image to an existing Word document."""
    try:
        document = load_document(doc_id)
        image_bytes = base64.b64decode(image_data)
        image_stream = BytesIO(image_bytes)
        document.add_picture(image_stream, width=Inches(width_inches))
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Image '{image_name}' added to document '{doc_id}.docx' successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding image: {str(e)}"

@mcp.tool()
def add_heading(doc_id: str, text: str, level: int, formatting: dict = None) -> str:
    """Adds a heading to an existing Word document with optional formatting.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        text (str): The heading text.
        level (int): The heading level (0-9, where 0 is Title).
        formatting (dict, optional): Dictionary with paragraph formatting options.
    """
    try:
        document = load_document(doc_id)
        
        # First, ensure the heading style exists in the document
        heading_style = None
        if level == 0:
            heading_style = "Title"
        else:
            heading_style = f"Heading {level}"
        
        # Check if this heading style exists in the document
        style_exists_in_doc = False
        for doc_style in document.styles:
            if doc_style.name == heading_style:
                style_exists_in_doc = True
                break
        
        # If style doesn't exist, it needs to be defined first
        if not style_exists_in_doc:
            try:
                # Add temporary paragraph to define the style
                temp_para = document.add_paragraph("", style=heading_style)
                # Remove the temporary paragraph
                p = temp_para._element
                p.getparent().remove(p)
            except KeyError:
                # If the style is not found, it's not a built-in style
                pass
        
        # Now add the actual heading
        heading = document.add_heading(text, level)
        
        # Apply formatting if provided
        if formatting:
            _apply_paragraph_formatting(heading, formatting)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return "Heading added successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding heading: {str(e)}"

# Table operations
@mcp.tool()
def add_table(doc_id: str, rows: int, cols: int, data: str = None, style: str = None) -> str:
    """Adds a table to an existing Word document.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        rows (int): The number of rows in the table.
        cols (int): The number of columns in the table.
        data (str, optional): Comma-separated data for the table cells (row-wise).
        style (str, optional): Table style (e.g., "Table Grid", "Light Shading").
    """
    try:
        document = load_document(doc_id)
        
        # If style is specified, ensure it exists in the document
        if style:
            style_exists_in_doc = False
            for doc_style in document.styles:
                if doc_style.name == style and doc_style.type == WD_STYLE_TYPE.TABLE:
                    style_exists_in_doc = True
                    break
            
            # If style doesn't exist, it might be a built-in style that needs defining
            if not style_exists_in_doc:
                try:
                    # Add temporary table to define the style
                    temp_table = document.add_table(rows=1, cols=1)
                    temp_table.style = style
                    # Remove the temporary table
                    p = temp_table._element
                    p.getparent().remove(p)
                except KeyError:
                    return f"Warning: Table style '{style}' not found. Table will be added with default style."
        
        # Create table with specified dimensions
        table = document.add_table(rows=rows, cols=cols)
        
        # Apply style if specified
        if style:
            try:
                table.style = style
            except KeyError:
                return f"Warning: Style '{style}' not found. Table added with default style."
        
        # Fill with data if provided
        if data:
            data_list = data.split(',')
            
            # Check if data matches table dimensions
            if len(data_list) > rows * cols:
                return f"Error: Number of data elements ({len(data_list)}) exceeds table dimensions ({rows}x{cols})."
                
            # Pad with empty strings if too few data elements
            if len(data_list) < rows * cols:
                data_list.extend([''] * (rows * cols - len(data_list)))
                
            # Fill table cells
            for i in range(rows):
                for j in range(cols):
                    cell_idx = i * cols + j
                    if cell_idx < len(data_list):
                        table.cell(i, j).text = data_list[cell_idx].strip()
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Table with {rows} rows and {cols} columns added successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding table: {str(e)}"

@mcp.tool()
def merge_table_cells(doc_id: str, table_index: int, start_row: int, start_col: int, end_row: int, end_col: int) -> str:
    """Merges cells in a table.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        table_index (int): The index of the table in the document (0-based).
        start_row (int): The starting row index (0-based).
        start_col (int): The starting column index (0-based).
        end_row (int): The ending row index (0-based).
        end_col (int): The ending column index (0-based).
    """
    try:
        document = load_document(doc_id)
        
        # Check if document has tables
        if not document.tables or len(document.tables) <= table_index:
            return f"Error: Table index {table_index} is out of range. Document has {len(document.tables) if document.tables else 0} tables."
        
        table = document.tables[table_index]
        
        # Validate row and column indices
        if start_row < 0 or end_row >= len(table.rows) or start_col < 0:
            return "Error: Row or column index out of range."
        
        # Get first cell in the range
        first_cell = table.cell(start_row, start_col)
        
        # Get last cell in the range
        last_cell = table.cell(end_row, end_col)
        
        # Merge the cells
        first_cell.merge(last_cell)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Cells merged from ({start_row},{start_col}) to ({end_row},{end_col}) in table {table_index}."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error merging table cells: {str(e)}"

@mcp.tool()
def get_table_data(doc_id: str, table_index: int, include_empty_cells: bool = True) -> str:
    """Retrieves a table's data as formatted text.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        table_index (int): The index of the table in the document (0-based).
        include_empty_cells (bool): Whether to include empty cells in the output.
    """
    try:
        document = load_document(doc_id)
        
        # Check if document has tables
        if not document.tables or len(document.tables) <= table_index:
            return f"Error: Table index {table_index} is out of range. Document has {len(document.tables) if document.tables else 0} tables."
        
        table = document.tables[table_index]
        
        # Get table data
        result = []
        for i, row in enumerate(table.rows):
            row_data = []
            
            # Handle cells before the first actual cell
            for _ in range(row.grid_cols_before):
                if include_empty_cells:
                    row_data.append("")
            
            # Handle actual cells
            for cell in row.cells:
                # Get all text from the cell, including from nested tables
                cell_text = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip() or include_empty_cells:
                        cell_text.append(paragraph.text)
                
                # Add nested tables as a note
                if cell.tables:
                    cell_text.append("[Contains nested table]")
                
                row_data.append("\n".join(cell_text))
            
            # Handle cells after the last actual cell
            for _ in range(row.grid_cols_after):
                if include_empty_cells:
                    row_data.append("")
            
            result.append(" | ".join(row_data))
        
        return "\n".join(result)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error retrieving table data: {str(e)}"

@mcp.tool()
def list_tables(doc_id: str) -> str:
    """Lists all tables in a document with their dimensions.
    
    Args:
        doc_id (str): The document ID (filename without extension).
    """
    try:
        document = load_document(doc_id)
        
        if not document.tables:
            return f"No tables found in document '{doc_id}.docx'."
        
        tables_info = []
        for i, table in enumerate(document.tables):
            row_count = len(table.rows)
            # Get the maximum number of columns across all rows
            col_count = max([row.grid_cols_before + len(row.cells) + row.grid_cols_after 
                            for row in table.rows]) if row_count > 0 else 0
            
            first_cell_text = ""
            if row_count > 0 and len(table.rows[0].cells) > 0:
                first_cell_text = table.rows[0].cells[0].text[:30]
                if len(table.rows[0].cells[0].text) > 30:
                    first_cell_text += "..."
            
            # Get table style if available
            style_name = "Default"
            if hasattr(table, "style") and table.style:
                style_name = table.style.name
            
            tables_info.append(f"Table {i}: {row_count} rows x {col_count} columns. Style: '{style_name}'. First cell: '{first_cell_text}'")
        
        return "\n".join(tables_info)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error listing tables: {str(e)}"

# Text formatting operations
@mcp.tool()
def set_paragraph_properties(doc_id: str, paragraph_index: int, formatting: dict) -> str:
    """Sets various properties of a paragraph.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        paragraph_index (int): Index of the paragraph to format (0-based).
        formatting (dict): Dictionary with paragraph formatting options like:
            - alignment: 'LEFT', 'CENTER', 'RIGHT', 'JUSTIFY'
            - left_indent, right_indent, first_line_indent: Indentation in inches
            - space_before, space_after: Spacing in points
            - line_spacing: Line spacing as multiple or points
            - keep_together, keep_with_next, page_break_before, widow_control: Boolean pagination options
    """
    try:
        document = load_document(doc_id)
        
        if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
            return "Error: Paragraph index out of range."
        
        paragraph = document.paragraphs[paragraph_index]
        _apply_paragraph_formatting(paragraph, formatting)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Paragraph {paragraph_index} properties set successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error setting paragraph properties: {str(e)}"

@mcp.tool()
def set_text_properties(doc_id: str, paragraph_index: int, run_index: int, formatting: dict) -> str:
    """Sets various properties of a run of text.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        paragraph_index (int): Index of the paragraph containing the run (0-based).
        run_index (int): Index of the run within the paragraph (0-based).
        formatting (dict): Dictionary with text formatting options like:
            - name: Font name
            - size: Font size in points
            - bold, italic, underline: Boolean style options
            - color: Color as hex (#RRGGBB) or rgb(r,g,b)
    """
    try:
        document = load_document(doc_id)
        
        if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
            return "Error: Paragraph index out of range."
        
        paragraph = document.paragraphs[paragraph_index]
        
        if run_index < 0 or run_index >= len(paragraph.runs):
            return f"Error: Run index {run_index} is out of range. Paragraph has {len(paragraph.runs)} runs."
        
        run = paragraph.runs[run_index]
        _apply_run_formatting(run, formatting)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Text properties set for run {run_index} in paragraph {paragraph_index}."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error setting text properties: {str(e)}"

# Utility functions
@mcp.tool()
def list_styles(doc_id: str) -> str:
    """Lists available paragraph and character styles in the document."""
    try:
        document = load_document(doc_id)
        para_styles = []
        char_styles = []
        table_styles = []
        
        for style in document.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                para_styles.append(style.name)
            elif style.type == WD_STYLE_TYPE.CHARACTER:
                char_styles.append(style.name)
            elif style.type == WD_STYLE_TYPE.TABLE:
                table_styles.append(style.name)
        
        result = []
        if para_styles:
            result.append("Paragraph styles:\n" + ", ".join(para_styles))
        if char_styles:
            result.append("\nCharacter styles:\n" + ", ".join(char_styles))
        if table_styles:
            result.append("\nTable styles:\n" + ", ".join(table_styles))
            
        return "\n".join(result)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error listing styles: {str(e)}"

@mcp.tool()
def add_header(doc_id: str, section_index: int, text: str = None, content: list = None) -> str:
    """Adds or modifies a header for a specific section.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        section_index (int): The index of the section to add a header to (0-based).
        text (str, optional): Simple text to add to the header.
        content (list, optional): Complex content for the header, following the same format
                                 as document content in create_complete_document.
                                 
    Returns:
        str: A message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        if not document.sections or section_index >= len(document.sections):
            return f"Error: Section index {section_index} is out of range. Document has {len(document.sections) if document.sections else 0} sections."
        
        section = document.sections[section_index]
        header = section.header
        
        # Unlink from previous if it's currently linked
        if header.is_linked_to_previous:
            header.is_linked_to_previous = False
        
        # Clear existing content
        for paragraph in header.paragraphs[1:]:
            p = paragraph._element
            p.getparent().remove(p)
        
        # If first paragraph exists, use it, otherwise add one
        if header.paragraphs:
            first_paragraph = header.paragraphs[0]
            if text:
                first_paragraph.text = text
            else:
                first_paragraph.text = ""
        else:
            if text:
                header.paragraphs[0].text = text
        
        # If complex content is provided, add it
        if content:
            for item in content:
                content_type = item.get("type", "").lower()
                item_text = item.get("text", "")
                
                if content_type == "paragraph":
                    style = item.get("style", "Header")
                    
                    # Check if style exists
                    style_exists_in_doc = False
                    for doc_style in document.styles:
                        if doc_style.name == style:
                            style_exists_in_doc = True
                            break
                    
                    # If style doesn't exist, try to define it
                    if not style_exists_in_doc and style == "Header":
                        try:
                            # Add temporary paragraph to define the style
                            temp_para = document.add_paragraph("", style=style)
                            # Remove the temporary paragraph
                            p = temp_para._element
                            p.getparent().remove(p)
                        except KeyError:
                            style = None  # Style not found
                    
                    # Add the paragraph
                    para = header.add_paragraph(item_text)
                    if style:
                        try:
                            para.style = style
                        except:
                            pass  # Style not found, continue with default
                    
                    # Apply formatting if provided
                    formatting = item.get("formatting", {})
                    if formatting:
                        _apply_paragraph_formatting(para, formatting)
                
                elif content_type == "table":
                    rows = item.get("rows", 1)
                    cols = item.get("cols", 1)
                    data = item.get("data", "")
                    
                    table = header.add_table(rows=rows, cols=cols)
                    
                    # Fill with data if provided
                    if data:
                        data_list = data.split(',')
                        
                        # Pad with empty strings if too few data elements
                        if len(data_list) < rows * cols:
                            data_list.extend([''] * (rows * cols - len(data_list)))
                            
                        # Fill table cells
                        for i in range(rows):
                            for j in range(cols):
                                cell_idx = i * cols + j
                                if cell_idx < len(data_list):
                                    table.cell(i, j).text = data_list[cell_idx].strip()
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Header added/modified for section {section_index}."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding header: {str(e)}"

@mcp.tool()
def add_footer(doc_id: str, section_index: int, text: str = None, content: list = None) -> str:
    """Adds or modifies a footer for a specific section.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        section_index (int): The index of the section to add a footer to (0-based).
        text (str, optional): Simple text to add to the footer.
        content (list, optional): Complex content for the footer, following the same format
                                 as document content in create_complete_document.
                                 
    Returns:
        str: A message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        if not document.sections or section_index >= len(document.sections):
            return f"Error: Section index {section_index} is out of range. Document has {len(document.sections) if document.sections else 0} sections."
        
        section = document.sections[section_index]
        footer = section.footer
        
        # Unlink from previous if it's currently linked
        if footer.is_linked_to_previous:
            footer.is_linked_to_previous = False
        
        # Clear existing content
        for paragraph in footer.paragraphs[1:]:
            p = paragraph._element
            p.getparent().remove(p)
        
        # If first paragraph exists, use it, otherwise add one
        if footer.paragraphs:
            first_paragraph = footer.paragraphs[0]
            if text:
                first_paragraph.text = text
            else:
                first_paragraph.text = ""
        else:
            if text:
                footer.paragraphs[0].text = text
        
        # If complex content is provided, add it
        if content:
            for item in content:
                content_type = item.get("type", "").lower()
                item_text = item.get("text", "")
                
                if content_type == "paragraph":
                    style = item.get("style", "Footer")
                    
                    # Check if style exists
                    style_exists_in_doc = False
                    for doc_style in document.styles:
                        if doc_style.name == style:
                            style_exists_in_doc = True
                            break
                    
                    # If style doesn't exist, try to define it
                    if not style_exists_in_doc and style == "Footer":
                        try:
                            # Add temporary paragraph to define the style
                            temp_para = document.add_paragraph("", style=style)
                            # Remove the temporary paragraph
                            p = temp_para._element
                            p.getparent().remove(p)
                        except KeyError:
                            style = None  # Style not found
                    
                    # Add the paragraph
                    para = footer.add_paragraph(item_text)
                    if style:
                        try:
                            para.style = style
                        except:
                            pass  # Style not found, continue with default
                    
                    # Apply formatting if provided
                    formatting = item.get("formatting", {})
                    if formatting:
                        _apply_paragraph_formatting(para, formatting)
                
                elif content_type == "table":
                    rows = item.get("rows", 1)
                    cols = item.get("cols", 1)
                    data = item.get("data", "")
                    
                    table = footer.add_table(rows=rows, cols=cols)
                    
                    # Fill with data if provided
                    if data:
                        data_list = data.split(',')
                        
                        # Pad with empty strings if too few data elements
                        if len(data_list) < rows * cols:
                            data_list.extend([''] * (rows * cols - len(data_list)))
                            
                        # Fill table cells
                        for i in range(rows):
                            for j in range(cols):
                                cell_idx = i * cols + j
                                if cell_idx < len(data_list):
                                    table.cell(i, j).text = data_list[cell_idx].strip()
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Footer added/modified for section {section_index}."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding footer: {str(e)}"

@mcp.tool()
def add_zoned_header(doc_id: str, section_index: int, left_text: str = "", center_text: str = "", right_text: str = "") -> str:
    """Adds a three-zone header with left, center, and right aligned text.
    
    This is a convenience function that creates a properly formatted header
    with text aligned in the left, center, and right zones using tab stops.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        section_index (int): The index of the section to add a header to (0-based).
        left_text (str): Text for the left zone.
        center_text (str): Text for the center zone.
        right_text (str): Text for the right zone.
        
    Returns:
        str: A message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        if not document.sections or section_index >= len(document.sections):
            return f"Error: Section index {section_index} is out of range. Document has {len(document.sections) if document.sections else 0} sections."
        
        section = document.sections[section_index]
        header = section.header
        
        # Unlink from previous if it's currently linked
        if header.is_linked_to_previous:
            header.is_linked_to_previous = False
        
        # Create the zoned header with tab-separated text
        header_text = f"{left_text}"
        if center_text:
            header_text += f"\t{center_text}"
            if right_text:
                header_text += f"\t{right_text}"
        elif right_text:
            header_text += f"\t\t{right_text}"
        
        # Clear existing content
        for paragraph in header.paragraphs[1:]:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Check if "Header" style exists and define it if needed
        header_style_exists = False
        for style in document.styles:
            if style.name == "Header":
                header_style_exists = True
                break
        
        if not header_style_exists:
            try:
                # Add temporary paragraph to define the style
                temp_para = document.add_paragraph("", style="Header")
                # Remove the temporary paragraph
                p = temp_para._element
                p.getparent().remove(p)
            except KeyError:
                pass  # Style not found, continue with default style
        
        # Apply the text to the first paragraph
        if header.paragraphs:
            paragraph = header.paragraphs[0]
            paragraph.text = header_text
            try:
                paragraph.style = document.styles["Header"]
            except:
                pass  # Style not found, continue with default style
        else:
            paragraph = header.add_paragraph(header_text)
            try:
                paragraph.style = document.styles["Header"]
            except:
                pass
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Zoned header added for section {section_index}."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding zoned header: {str(e)}"

@mcp.tool()
def add_zoned_footer(doc_id: str, section_index: int, left_text: str = "", center_text: str = "", right_text: str = "") -> str:
    """Adds a three-zone footer with left, center, and right aligned text.
    
    This is a convenience function that creates a properly formatted footer
    with text aligned in the left, center, and right zones using tab stops.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        section_index (int): The index of the section to add a footer to (0-based).
        left_text (str): Text for the left zone.
        center_text (str): Text for the center zone.
        right_text (str): Text for the right zone.
        
    Returns:
        str: A message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        if not document.sections or section_index >= len(document.sections):
            return f"Error: Section index {section_index} is out of range. Document has {len(document.sections) if document.sections else 0} sections."
        
        section = document.sections[section_index]
        footer = section.footer
        
        # Unlink from previous if it's currently linked
        if footer.is_linked_to_previous:
            footer.is_linked_to_previous = False
        
        # Create the zoned footer with tab-separated text
        footer_text = f"{left_text}"
        if center_text:
            footer_text += f"\t{center_text}"
            if right_text:
                footer_text += f"\t{right_text}"
        elif right_text:
            footer_text += f"\t\t{right_text}"
        
        # Clear existing content
        for paragraph in footer.paragraphs[1:]:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Check if "Footer" style exists and define it if needed
        footer_style_exists = False
        for style in document.styles:
            if style.name == "Footer":
                footer_style_exists = True
                break
        
        if not footer_style_exists:
            try:
                # Add temporary paragraph to define the style
                temp_para = document.add_paragraph("", style="Footer")
                # Remove the temporary paragraph
                p = temp_para._element
                p.getparent().remove(p)
            except KeyError:
                pass  # Style not found, continue with default style
        
        # Apply the text to the first paragraph
        if footer.paragraphs:
            paragraph = footer.paragraphs[0]
            paragraph.text = footer_text
            try:
                paragraph.style = document.styles["Footer"]
            except:
                pass  # Style not found, continue with default style
        else:
            paragraph = footer.add_paragraph(footer_text)
            try:
                paragraph.style = document.styles["Footer"]
            except:
                pass
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Zoned footer added for section {section_index}."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding zoned footer: {str(e)}"

@mcp.tool()
def remove_header(doc_id: str, section_index: int) -> str:
    """Removes the header from a specific section, linking it to the previous section.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        section_index (int): The index of the section to remove the header from (0-based).
        
    Returns:
        str: A message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        if not document.sections or section_index >= len(document.sections):
            return f"Error: Section index {section_index} is out of range. Document has {len(document.sections) if document.sections else 0} sections."
        
        section = document.sections[section_index]
        header = section.header
        
        # Link to previous, which removes this header definition
        header.is_linked_to_previous = True
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Header removed from section {section_index}."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error removing header: {str(e)}"

@mcp.tool()
def remove_footer(doc_id: str, section_index: int) -> str:
    """Removes the footer from a specific section, linking it to the previous section.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        section_index (int): The index of the section to remove the footer from (0-based).
        
    Returns:
        str: A message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        if not document.sections or section_index >= len(document.sections):
            return f"Error: Section index {section_index} is out of range. Document has {len(document.sections) if document.sections else 0} sections."
        
        section = document.sections[section_index]
        footer = section.footer
        
        # Link to previous, which removes this footer definition
        footer.is_linked_to_previous = True
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Footer removed from section {section_index}."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error removing footer: {str(e)}"

@mcp.tool()
def get_header_text(doc_id: str, section_index: int) -> str:
    """Gets the text content of a header for a specific section.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        section_index (int): The index of the section (0-based).
        
    Returns:
        str: The text content of the header or status message.
    """
    try:
        document = load_document(doc_id)
        
        if not document.sections or section_index >= len(document.sections):
            return f"Error: Section index {section_index} is out of range. Document has {len(document.sections) if document.sections else 0} sections."
        
        section = document.sections[section_index]
        header = section.header
        
        if header.is_linked_to_previous:
            # Find the first previous section with a header definition
            linked_section_index = section_index
            while linked_section_index > 0:
                linked_section_index -= 1
                prev_header = document.sections[linked_section_index].header
                if not prev_header.is_linked_to_previous:
                    return f"Header is linked to section {linked_section_index}. Content: {get_header_text(doc_id, linked_section_index)}"
            
            return "No header defined for this section (linked to previous, but no previous header found)."
        
        # Header has its own definition, extract the text
        header_text = []
        for paragraph in header.paragraphs:
            header_text.append(paragraph.text)
        
        if not header_text:
            return "Header is defined but contains no text."
        
        return "\n".join(header_text)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error getting header text: {str(e)}"

@mcp.tool()
def get_footer_text(doc_id: str, section_index: int) -> str:
    """Gets the text content of a footer for a specific section.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        section_index (int): The index of the section (0-based).
        
    Returns:
        str: The text content of the footer or status message.
    """
    try:
        document = load_document(doc_id)
        
        if not document.sections or section_index >= len(document.sections):
            return f"Error: Section index {section_index} is out of range. Document has {len(document.sections) if document.sections else 0} sections."
        
        section = document.sections[section_index]
        footer = section.footer
        
        if footer.is_linked_to_previous:
            # Find the first previous section with a footer definition
            linked_section_index = section_index
            while linked_section_index > 0:
                linked_section_index -= 1
                prev_footer = document.sections[linked_section_index].footer
                if not prev_footer.is_linked_to_previous:
                    return f"Footer is linked to section {linked_section_index}. Content: {get_footer_text(doc_id, linked_section_index)}"
            
            return "No footer defined for this section (linked to previous, but no previous footer found)."
        
        # Footer has its own definition, extract the text
        footer_text = []
        for paragraph in footer.paragraphs:
            footer_text.append(paragraph.text)
        
        if not footer_text:
            return "Footer is defined but contains no text."
        
        return "\n".join(footer_text)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error getting footer text: {str(e)}"

@mcp.tool()
def convert_to_pdf(doc_id: str) -> str:
   """Converts a Word document to PDF format."""
   try:
       doc_path = get_document_path(doc_id)
       if not os.path.exists(doc_path):
           return f"Error: Document '{doc_id}.docx' not found."
       
       script_dir = os.path.dirname(os.path.abspath(__file__))
       pdf_path = os.path.join(script_dir, f"{doc_id}.pdf")
       
       convert(doc_path, pdf_path)
       
       return f"Document successfully converted to PDF at: {os.path.abspath(pdf_path)}"
   except Exception as e:
       return f"Error converting document to PDF: {str(e)}"

@mcp.tool()
def analyze_document_structure(doc_id: str) -> str:
   """Analyzes the structure of a document, showing paragraphs, runs, and tables.
   
   This tool provides a detailed breakdown of the document structure,
   which can be helpful for understanding how to target specific elements
   for modification.
   
   Args:
       doc_id (str): The document ID (filename without extension).
   """
   try:
       document = load_document(doc_id)
       
       structure = []
       structure.append(f"Document Structure Analysis for '{doc_id}.docx':")
       structure.append(f"Total paragraphs: {len(document.paragraphs)}")
       structure.append(f"Total tables: {len(document.tables)}")
       structure.append("\nParagraph Details:")
       
       for i, para in enumerate(document.paragraphs):
           if not para.text.strip():
               structure.append(f"  Paragraph {i}: [Empty paragraph]")
               continue
               
           style = para.style.name if para.style else "Default"
           run_count = len(para.runs)
           structure.append(f"  Paragraph {i}: Style='{style}', Runs={run_count}")
           structure.append(f"    Text: \"{para.text[:50]}{'...' if len(para.text) > 50 else ''}\"")
           
           if run_count > 0:
               structure.append(f"    Run details:")
               for j, run in enumerate(para.runs):
                   bold = "Bold" if run.bold else "Normal"
                   italic = "Italic" if run.italic else "Normal"
                   style_name = run.style.name if run.style else "Default"
                   structure.append(f"      Run {j}: Style='{style_name}', {bold}, {italic}, Text=\"{run.text[:30]}{'...' if len(run.text) > 30 else ''}\"")
       
       if document.tables:
           structure.append("\nTable Details:")
           for i, table in enumerate(document.tables):
               row_count = len(table.rows)
               col_count = max([row.grid_cols_before + len(row.cells) + row.grid_cols_after 
                               for row in table.rows]) if row_count > 0 else 0
               
               style_name = table.style.name if table.style else "Default"
               
               structure.append(f"  Table {i}: {row_count} rows x {col_count} columns")
               structure.append(f"    Style: {style_name}")
               
               # Show a preview of the first few cells
               if row_count > 0 and len(table.rows[0].cells) > 0:
                   structure.append(f"    Preview:")
                   max_preview_rows = min(3, row_count)
                   for r in range(max_preview_rows):
                       row = table.rows[r]
                       cell_texts = []
                       for cell in row.cells[:min(3, len(row.cells))]:
                           cell_text = cell.text[:20]
                           if len(cell.text) > 20:
                               cell_text += "..."
                           cell_texts.append(f"\"{cell_text}\"")
                       
                       additional = "..." if len(row.cells) > 3 else ""
                       structure.append(f"      Row {r}: {', '.join(cell_texts)}{additional}")
       
       return "\n".join(structure)
   except ValueError as e:
       return str(e)
   except Exception as e:
       return f"Error analyzing document structure: {str(e)}"

@mcp.prompt()
def word_document_usage() -> str:
    """Provides guidance on how to use this MCP server with Word documents."""
    return """
# Word Document Server Usage Guide

This server allows you to create, read, and manipulate Microsoft Word (.docx) documents. Here's how to use it effectively:

## Reading Documents
To read an existing document:
- Use the resource: `word://document_name/content` (replace "document_name" with the filename without .docx extension)
- Or call the tool: `read_document("document_name")`

Example: To read a file named "bitcoin_overview.docx":
- Request the resource: `word://bitcoin_overview/content`
- Or call: `read_document("bitcoin_overview")`

## Working with Styles
Word documents heavily rely on styles for consistent formatting. This server provides several tools for working with styles:

### Style Management
- `ensure_style_exists("document_name", "Heading 1", "paragraph")` - Ensures a built-in style exists in the document
- `create_custom_style("document_name", "MyCustomStyle", "paragraph", "Normal")` - Creates a new custom style
- `modify_style("document_name", "MyStyle", {"font": {"size": 12, "bold": True}, "paragraph": {"alignment": "CENTER"}})` - Modifies a style
- `get_styles_detail("document_name", "paragraph")` - Gets detailed information about styles
- `check_style_usage("document_name", "Heading 1")` - Checks where a style is used in a document
- `list_styles("document_name")` - Lists all styles in a document

### Style Tips
- Always ensure a style exists before using it (Word ignores undefined styles)
- Use `ensure_style_exists()` to define built-in styles before using them
- Apply styles by name when adding content: `add_paragraph("document_name", "Text", style="Heading 1")`
- Custom styles can be based on existing styles using the `base_style` parameter
- Style changes affect all content using that style

## Creating Documents

Step-by-Step Method

First create a document: create_document("my_doc", "My Document Title")
Add content using any of these tools:

add_paragraph("my_doc", "This is a paragraph of text", style="Normal", formatting={"alignment": "CENTER"})
add_heading("my_doc", "Section Heading", 1) (levels 0-4, where 0 is title)
add_table("my_doc", 3, 3, "Cell 1,Cell 2,Cell 3,Cell 4,Cell 5,Cell 6,Cell 7,Cell 8,Cell 9", "Table Grid")
add_image("my_doc", base64_image_data, "image.png", 4.0)

Updating Documents

Append content to existing documents:
pythonCopyappend_to_document("my_doc", [{"type": "paragraph", "text": "New content", "style": "Normal"}])

Replace existing document:
pythonCopyreplace_document("my_doc", "New Title", [{"type": "paragraph", "text": "Replacement content", "style": "Normal"}])


Text Formatting
Paragraph Formatting
Set paragraph properties:
pythonCopyset_paragraph_properties("my_doc", 1, {
    "alignment": "CENTER",
    "left_indent": 0.5,        # in inches
    "right_indent": 0.5,       # in inches
    "first_line_indent": 0.25, # in inches
    "space_before": 12,        # in points
    "space_after": 12,         # in points
    "line_spacing": 1.5,       # multiple or points
    "keep_together": True,     # pagination options"
    "keep_with_next": True,
    "page_break_before": False,
    "widow_control": True
})

Text/Run Formatting
Add formatted text to an existing paragraph:
pythonCopyadd_formatted_text("my_doc", 1, "This text will be formatted", {
    "name": "Arial",          # font name
    "size": 12,               # point size
    "bold": True,             # Boolean
    "italic": False,          # Boolean
    "underline": True,        # Boolean
    "color": "#FF0000"        # hex color or rgb(r,g,b)
})
Or set properties of an existing text run:
pythonCopyset_text_properties("my_doc", 1, 0, {  # paragraph 1, run 0
    "bold": True,
    "color": "rgb(0,0,255)"
})
Table Operations

Create tables: add_table("my_doc", 3, 3, "A,B,C,D,E,F,G,H,I", "Table Grid")
Merge cells: merge_table_cells("my_doc", 0, 0, 0, 0, 1) (table 0, from cell (0,0) to (0,1))
Get table data: get_table_data("my_doc", 0)
List tables: list_tables("my_doc")

Section Operations
Sections define page layout settings like margins and orientation:
Working with Sections

Add a section: add_section("my_doc", "NEW_PAGE")  # Other options: "EVEN_PAGE", "ODD_PAGE", "CONTINUOUS"
List sections: list_sections("my_doc")
Change page orientation: change_page_orientation("my_doc", 0, "LANDSCAPE")  # Change section 0 to landscape

Set Section Properties
pythonCopyset_section_properties("my_doc", 0, {  # section 0
    "orientation": "LANDSCAPE",
    "page_width": 11,         # in inches
    "page_height": 8.5,       # in inches
    "left_margin": 1,         # in inches
    "right_margin": 1,        # in inches
    "top_margin": 0.75,       # in inches
    "bottom_margin": 0.75,    # in inches
    "header_distance": 0.5,   # in inches
    "footer_distance": 0.5    # in inches
})
Headers and Footers
Headers and footers are linked to sections and provide content that appears at the top/bottom of each page:
Basic Headers and Footers

Add a simple header: add_header("my_doc", 0, "My Document Title")  # Add to section 0
Add a simple footer: add_footer("my_doc", 0, "Page X of Y")  # Add to section 0
Remove a header: remove_header("my_doc", 0)  # Remove from section 0
Remove a footer: remove_footer("my_doc", 0)  # Remove from section 0

Three-Zone Headers and Footers
Create headers/footers with left, center, and right aligned content:
pythonCopyadd_zoned_header("my_doc", 0, 
                "Left Text",      # Left-aligned
                "Center Text",    # Center-aligned
                "Right Text")     # Right-aligned

add_zoned_footer("my_doc", 0, 
                "Author: John Doe",  # Left-aligned
                "Confidential",      # Center-aligned
                "Page 1")            # Right-aligned
Getting Header/Footer Content

Get header text: get_header_text("my_doc", 0)  # From section 0
Get footer text: get_footer_text("my_doc", 0)  # From section 0

Analyzing Documents
To understand a document's structure before modifying it:

analyze_document_structure("my_doc")

Converting to PDF
To convert a Word document to PDF format:

convert_to_pdf("my_document") - This will create a PDF with the same name in the server directory

Utility Functions

Check if a document exists: check_document_exists("my_document")
List all available documents: list_available_documents()
List available styles in a document: list_styles("my_document")

Tips for Working with Word Documents

Style Management:

Check if a style exists before using it with ensure_style_exists()
Word ignores style applications if the style isn't defined in the document
Built-in styles need to be defined (used at least once) before they appear in the document
Modifying a style affects all content using that style


Document Structure:

Use analyze_document_structure() to understand document contents
Check paragraph indexes carefully (they start at 0)
Remember that changes are saved immediately


Working with Sections:

When changing orientation, you may need to explicitly set page dimensions
Headers and footers are linked to sections, so create sections first if needed
Most documents start with one section by default



Remember to check the document path returned by create_document() to know where your files are stored.
"""

if __name__ == "__main__":
    mcp.run()
