"""
Document Headers Tools

This module provides tools for working with headers and footers in Word documents.
"""

from docx import Document
import os
from typing import Dict, Optional, Any, List

def register_document_headers_tools(mcp):
    """
    Registers all document header and footer-related tools with the MCP server.
    
    Args:
        mcp: The MCP server instance to register tools with
    """
    
    @mcp.tool()
    def add_header(document_name: str, section_index: int, text: str) -> str:
        """
        Adds a header to a specific section of the document.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section to add the header to (0-based)
            text: The text to add to the header
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            header = section.header
            header.paragraphs[0].text = text
            
            doc.save(doc_path)
            return f"Header added to section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error adding header: {str(e)}"
    
    @mcp.tool()
    def add_footer(document_name: str, section_index: int, text: str) -> str:
        """
        Adds a footer to a specific section of the document.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section to add the footer to (0-based)
            text: The text to add to the footer
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            footer = section.footer
            footer.paragraphs[0].text = text
            
            doc.save(doc_path)
            return f"Footer added to section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error adding footer: {str(e)}"
    
    @mcp.tool()
    def remove_header(document_name: str, section_index: int) -> str:
        """
        Removes the header from a specific section of the document.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section to remove the header from (0-based)
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            section.header.is_linked_to_previous = True
            
            doc.save(doc_path)
            return f"Header removed from section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error removing header: {str(e)}"
    
    @mcp.tool()
    def remove_footer(document_name: str, section_index: int) -> str:
        """
        Removes the footer from a specific section of the document.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section to remove the footer from (0-based)
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            section.footer.is_linked_to_previous = True
            
            doc.save(doc_path)
            return f"Footer removed from section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error removing footer: {str(e)}"
    
    @mcp.tool()
    def add_zoned_header(document_name: str, section_index: int, 
                        left_text: str, center_text: str, right_text: str) -> str:
        """
        Adds a header with left, center, and right aligned text to a specific section.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section to add the header to (0-based)
            left_text: Text to align on the left
            center_text: Text to align in the center
            right_text: Text to align on the right
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            header = section.header
            
            # Clear existing content
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    run.text = ""
            
            # Create a table for the zoned header
            table = header.add_table(1, 3)
            table.cell(0, 0).text = left_text
            table.cell(0, 1).text = center_text
            table.cell(0, 2).text = right_text
            
            # Set column widths and alignment
            table.cell(0, 0).paragraphs[0].alignment = 0  # LEFT
            table.cell(0, 1).paragraphs[0].alignment = 1  # CENTER
            table.cell(0, 2).paragraphs[0].alignment = 2  # RIGHT
            
            # Hide table borders
            for cell in table._cells:
                for border in cell._element.xpath('.//w:tcBorders/w:*'):
                    border.set('w:val', 'none')
            
            doc.save(doc_path)
            return f"Zoned header added to section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error adding zoned header: {str(e)}"
    
    @mcp.tool()
    def add_zoned_footer(document_name: str, section_index: int, 
                        left_text: str, center_text: str, right_text: str) -> str:
        """
        Adds a footer with left, center, and right aligned text to a specific section.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section to add the footer to (0-based)
            left_text: Text to align on the left
            center_text: Text to align in the center
            right_text: Text to align on the right
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            footer = section.footer
            
            # Clear existing content
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    run.text = ""
            
            # Create a table for the zoned footer
            table = footer.add_table(1, 3)
            table.cell(0, 0).text = left_text
            table.cell(0, 1).text = center_text
            table.cell(0, 2).text = right_text
            
            # Set column widths and alignment
            table.cell(0, 0).paragraphs[0].alignment = 0  # LEFT
            table.cell(0, 1).paragraphs[0].alignment = 1  # CENTER
            table.cell(0, 2).paragraphs[0].alignment = 2  # RIGHT
            
            # Hide table borders
            for cell in table._cells:
                for border in cell._element.xpath('.//w:tcBorders/w:*'):
                    border.set('w:val', 'none')
            
            doc.save(doc_path)
            return f"Zoned footer added to section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error adding zoned footer: {str(e)}"
    
    @mcp.tool()
    def get_header_text(document_name: str, section_index: int) -> Dict[str, Any]:
        """
        Gets the text content of a header from a specific section.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section to get the header from (0-based)
            
        Returns:
            Header content information
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return {"error": f"Document {document_name}.docx does not exist"}
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return {"error": f"Section index {section_index} is out of range"}
                
            section = doc.sections[section_index]
            header = section.header
            
            content = []
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            return {
                "section_index": section_index,
                "content": content
            }
        except Exception as e:
            return {"error": f"Error getting header text: {str(e)}"}
    
    @mcp.tool()
    def get_footer_text(document_name: str, section_index: int) -> Dict[str, Any]:
        """
        Gets the text content of a footer from a specific section.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section to get the footer from (0-based)
            
        Returns:
            Footer content information
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return {"error": f"Document {document_name}.docx does not exist"}
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return {"error": f"Section index {section_index} is out of range"}
                
            section = doc.sections[section_index]
            footer = section.footer
            
            content = []
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            return {
                "section_index": section_index,
                "content": content
            }
        except Exception as e:
            return {"error": f"Error getting footer text: {str(e)}"}
    
    @mcp.tool()
    def set_footer_page_numbers(document_name: str, section_index: int, format_string: str = "Page {0} of {1}") -> str:
        """
        Adds page numbers to the footer of a specific section.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section to add page numbers to (0-based)
            format_string: Format string for page numbers. Use {0} for current page and {1} for total pages
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            footer = section.footer
            
            # Clear existing content
            for paragraph in footer.paragraphs:
                for run in paragraph.runs:
                    run.text = ""
            
            # Add a new paragraph if needed
            if len(footer.paragraphs) == 0:
                paragraph = footer.add_paragraph()
            else:
                paragraph = footer.paragraphs[0]
            
            # Set alignment to center
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add page number field codes
            paragraph.add_run(format_string.format("PAGE", "NUMPAGES"))
            
            # Convert runs to field codes
            xml = paragraph._p.xml
            xml = xml.replace("PAGE", "{PAGE}")
            xml = xml.replace("NUMPAGES", "{NUMPAGES}")
            paragraph._p._element.clear()
            paragraph._p._element.append(docx.oxml.parse_xml(xml))
            
            doc.save(doc_path)
            return f"Page numbers added to footer in section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error adding page numbers: {str(e)}"
    
    @mcp.tool()
    def format_footer(document_name: str, section_index: int, properties: Dict[str, Any]) -> str:
        """
        Sets formatting properties of a footer.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section containing the footer (0-based)
            properties: Dictionary of formatting properties for the footer paragraphs
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            footer = section.footer
            
            for paragraph in footer.paragraphs:
                # Apply paragraph formatting
                if 'alignment' in properties:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    alignment_map = {
                        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if properties['alignment'].upper() in alignment_map:
                        paragraph.alignment = alignment_map[properties['alignment'].upper()]
                
                # Apply font formatting to all runs
                if 'font' in properties:
                    for run in paragraph.runs:
                        if 'name' in properties['font']:
                            run.font.name = properties['font']['name']
                        if 'size' in properties['font']:
                            run.font.size = properties['font']['size']
                        if 'bold' in properties['font']:
                            run.font.bold = properties['font']['bold']
                        if 'italic' in properties['font']:
                            run.font.italic = properties['font']['italic']
                        if 'color' in properties['font']:
                            run.font.color.rgb = properties['font']['color']
            
            doc.save(doc_path)
            return f"Footer formatting applied in section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error formatting footer: {str(e)}"
    
    @mcp.tool()
    def add_footer_image(document_name: str, section_index: int, image_path: str, width: float = None, height: float = None) -> str:
        """
        Adds an image to the footer of a specific section.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section to add the image to (0-based)
            image_path: Path to the image file
            width: Optional width of the image in inches
            height: Optional height of the image in inches
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
            
            if not os.path.exists(image_path):
                return f"Error: Image file {image_path} does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            footer = section.footer
            
            # Add a new paragraph if needed
            if len(footer.paragraphs) == 0:
                paragraph = footer.add_paragraph()
            else:
                paragraph = footer.paragraphs[0]
            
            # Add the image
            run = paragraph.add_run()
            picture = run.add_picture(image_path)
            
            # Set width and height if provided
            if width is not None:
                from docx.shared import Inches
                picture.width = Inches(width)
            
            if height is not None:
                from docx.shared import Inches
                picture.height = Inches(height)
            
            doc.save(doc_path)
            return f"Image added to footer in section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error adding image to footer: {str(e)}"
    
    @mcp.tool()
    def list_footer_content(document_name: str, section_index: int = None) -> Dict[str, Any]:
        """
        Lists the content of footers in the document.
        
        Args:
            document_name: Name of the document without extension
            section_index: Optional index of a specific section to get footer content from
            
        Returns:
            Dictionary containing footer content information
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return {"error": f"Document {document_name}.docx does not exist"}
                
            doc = Document(doc_path)
            
            # Check specific section if provided
            if section_index is not None:
                if section_index < 0 or section_index >= len(doc.sections):
                    return {"error": f"Section index {section_index} is out of range"}
                
                sections_to_check = [doc.sections[section_index]]
            else:
                sections_to_check = doc.sections
            
            footer_info = []
            
            for i, section in enumerate(sections_to_check):
                if section_index is not None:
                    curr_index = section_index
                else:
                    curr_index = i
                
                footer = section.footer
                content = []
                
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text)
                
                has_tables = len(footer.tables) > 0
                
                footer_info.append({
                    "section_index": curr_index,
                    "content": content,
                    "has_tables": has_tables,
                    "linked_to_previous": footer.is_linked_to_previous
                })
            
            return {
                "footers": footer_info,
                "total_sections": len(doc.sections)
            }
        except Exception as e:
            return {"error": f"Error listing footer content: {str(e)}"}
    
    @mcp.tool()
    def set_different_first_page_footer(document_name: str, section_index: int, enable: bool = True) -> str:
        """
        Sets whether the first page of a section has a different footer.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section (0-based)
            enable: Whether to enable different first page footer
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            section.different_first_page_header_footer = enable
            
            doc.save(doc_path)
            return f"Different first page footer {'enabled' if enable else 'disabled'} in section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error setting different first page footer: {str(e)}"
    
    @mcp.tool()
    def get_footer_distance(document_name: str, section_index: int) -> Dict[str, Any]:
        """
        Gets the distance between the footer and the bottom of the page.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section (0-based)
            
        Returns:
            Dictionary with footer distance information
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return {"error": f"Document {document_name}.docx does not exist"}
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return {"error": f"Section index {section_index} is out of range"}
                
            section = doc.sections[section_index]
            footer_distance = section.footer_distance
            
            # Convert to inches from TWIPS (twentieth of a point)
            inches = footer_distance.inches
            
            return {
                "section_index": section_index,
                "footer_distance_inches": inches,
                "footer_distance_cm": inches * 2.54
            }
        except Exception as e:
            return {"error": f"Error getting footer distance: {str(e)}"}
    
    @mcp.tool()
    def set_footer_distance(document_name: str, section_index: int, distance_inches: float) -> str:
        """
        Sets the distance between the footer and the bottom of the page.
        
        Args:
            document_name: Name of the document without extension
            section_index: Index of the section (0-based)
            distance_inches: Distance in inches between footer and page bottom
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if section_index < 0 or section_index >= len(doc.sections):
                return f"Error: Section index {section_index} is out of range"
                
            section = doc.sections[section_index]
            from docx.shared import Inches
            section.footer_distance = Inches(distance_inches)
            
            doc.save(doc_path)
            return f"Footer distance set to {distance_inches} inches in section {section_index} in {document_name}.docx"
        except Exception as e:
            return f"Error setting footer distance: {str(e)}"

    print("Document headers tools registered")
    return mcp
