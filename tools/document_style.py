"""
Document Style Tools

This module provides tools for working with document styles in Word documents.
"""

from utils.helpers import load_document
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from typing import Dict, List, Any, Optional

def register_document_style_tools(mcp):
    """
    Registers all document style-related tools with the MCP server.
    
    Args:
        mcp: The MCP server instance to register tools with
    """
    
    @mcp.tool()
    def list_styles(document_name: str) -> Dict[str, Any]:
        """
        Lists all styles in a document.
        
        Args:
            document_name: Name of the document without extension
            
        Returns:
            Dictionary with style information
        """
        try:
            doc, doc_path = load_document(document_name)
            
            style_info = {
                "paragraph_styles": [],
                "character_styles": [],
                "table_styles": [],
                "list_styles": []
            }
            
            style_types = {
                WD_STYLE_TYPE.PARAGRAPH: "paragraph_styles",
                WD_STYLE_TYPE.CHARACTER: "character_styles",
                WD_STYLE_TYPE.TABLE: "table_styles",
                WD_STYLE_TYPE.LIST: "list_styles"
            }
            
            for style in doc.styles:
                if style.type in style_types:
                    style_category = style_types[style.type]
                    style_info[style_category].append(style.name)
            
            return style_info
        except FileNotFoundError as e:
            return {"error": str(e)}
        except Exception as e:
            return {"error": f"Error listing styles: {str(e)}"}
    
    @mcp.tool()
    def ensure_style_exists(document_name: str, style_name: str, style_type: str = "paragraph") -> str:
        """
        Ensures a built-in style exists in the document.
        
        Args:
            document_name: Name of the document without extension
            style_name: Name of the style to ensure exists
            style_type: Type of the style (paragraph, character, table, list)
            
        Returns:
            Message indicating success
        """
        try:
            doc, doc_path = load_document(document_name)
            
            # Attempt to use the style to ensure it's available
            if style_type.lower() == "paragraph":
                doc.add_paragraph("", style=style_name)
                # Remove the empty paragraph we just added
                doc._body._body.remove(doc._body._body[-1])
                save_document(doc, doc_path)
                return f"Style '{style_name}' is now available in {document_name}.docx"
            else:
                # For other style types, check if it exists
                for style in doc.styles:
                    if style.name == style_name:
                        return f"Style '{style_name}' already exists in {document_name}.docx"
                return f"Error: Could not ensure style '{style_name}' exists as it's not a built-in style"
        except FileNotFoundError as e:
            return str(e)
        except Exception as e:
            return f"Error ensuring style exists: {str(e)}"
    
    @mcp.tool()
    def create_custom_style(document_name: str, style_name: str, style_type: str = "paragraph", 
                           base_style: str = "Normal") -> str:
        """
        Creates a new custom style in the document.
        
        Args:
            document_name: Name of the document without extension
            style_name: Name of the style to create
            style_type: Type of the style (paragraph, character)
            base_style: Base style to inherit from
            
        Returns:
            Message indicating success
        """
        try:
            doc, doc_path = load_document(document_name)
            
            # Check if style already exists
            for style in doc.styles:
                if style.name == style_name:
                    return f"Style '{style_name}' already exists in {document_name}.docx"
            
            # Map style type string to enum
            style_type_map = {
                "paragraph": WD_STYLE_TYPE.PARAGRAPH,
                "character": WD_STYLE_TYPE.CHARACTER
            }
            
            if style_type.lower() not in style_type_map:
                return f"Error: Style type '{style_type}' is not supported. Use 'paragraph' or 'character'."
            
            # Create the new style based on the base style
            try:
                new_style = doc.styles.add_style(style_name, style_type_map[style_type.lower()])
                new_style.base_style = doc.styles[base_style]
                save_document(doc, doc_path)
                return f"Custom style '{style_name}' created in {document_name}.docx"
            except KeyError:
                return f"Error: Base style '{base_style}' not found in document"
        except FileNotFoundError as e:
            return str(e)
        except Exception as e:
            return f"Error creating custom style: {str(e)}"
    
    @mcp.tool()
    def modify_style(document_name: str, style_name: str, properties: Dict[str, Any]) -> str:
        """
        Modifies a style in the document.
        
        Args:
            document_name: Name of the document without extension
            style_name: Name of the style to modify
            properties: Dictionary of properties to modify
            
        Returns:
            Message indicating success
        """
        try:
            doc, doc_path = load_document(document_name)
            
            # Check if style exists
            try:
                style = doc.styles[style_name]
            except KeyError:
                return f"Error: Style '{style_name}' not found in document"
            
            # Modify font properties
            if "font" in properties:
                font_props = properties["font"]
                font = style.font
                
                if "name" in font_props:
                    font.name = font_props["name"]
                if "size" in font_props:
                    font.size = Pt(font_props["size"])
                if "bold" in font_props:
                    font.bold = font_props["bold"]
                if "italic" in font_props:
                    font.italic = font_props["italic"]
                if "underline" in font_props:
                    font.underline = font_props["underline"]
                if "color" in font_props:
                    color_str = parse_color(font_props["color"])
                    if color_str:
                        r = int(color_str[0:2], 16)
                        g = int(color_str[2:4], 16)
                        b = int(color_str[4:6], 16)
                        font.color.rgb = RGBColor(r, g, b)
            
            # Modify paragraph properties
            if "paragraph" in properties and hasattr(style, "paragraph_format"):
                para_props = properties["paragraph"]
                para_format = style.paragraph_format
                
                if "alignment" in para_props:
                    alignment_map = {
                        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if para_props["alignment"].upper() in alignment_map:
                        para_format.alignment = alignment_map[para_props["alignment"].upper()]
                
                if "left_indent" in para_props:
                    para_format.left_indent = Inches(para_props["left_indent"])
                if "right_indent" in para_props:
                    para_format.right_indent = Inches(para_props["right_indent"])
                if "first_line_indent" in para_props:
                    para_format.first_line_indent = Inches(para_props["first_line_indent"])
                if "space_before" in para_props:
                    para_format.space_before = Pt(para_props["space_before"])
                if "space_after" in para_props:
                    para_format.space_after = Pt(para_props["space_after"])
                if "line_spacing" in para_props:
                    para_format.line_spacing = para_props["line_spacing"]
            
            save_document(doc, doc_path)
            return f"Style '{style_name}' modified in {document_name}.docx"
        except FileNotFoundError as e:
            return str(e)
        except Exception as e:
            return f"Error modifying style: {str(e)}"
    
    @mcp.tool()
    def get_styles_detail(document_name: str, style_type: str = None) -> Dict[str, Any]:
        """
        Gets detailed information about styles in the document.
        
        Args:
            document_name: Name of the document without extension
            style_type: Optional type to filter by (paragraph, character, table, list)
            
        Returns:
            Dictionary with detailed style information
        """
        try:
            doc, _ = load_document(document_name)
            
            # Map style type string to enum if provided
            style_type_enum = None
            if style_type:
                style_type_map = {
                    "paragraph": WD_STYLE_TYPE.PARAGRAPH,
                    "character": WD_STYLE_TYPE.CHARACTER,
                    "table": WD_STYLE_TYPE.TABLE,
                    "list": WD_STYLE_TYPE.LIST
                }
                if style_type.lower() in style_type_map:
                    style_type_enum = style_type_map[style_type.lower()]
            
            styles_info = []
            
            for style in doc.styles:
                # Skip if not matching the requested type
                if style_type_enum and style.type != style_type_enum:
                    continue
                
                info = {
                    "name": style.name,
                    "type": str(style.type).split('.')[-1] if style.type else "UNKNOWN",
                    "built_in": style.built_in,
                    "base_style": style.base_style.name if style.base_style else None
                }
                
                # Get font details if available
                if hasattr(style, "font"):
                    info["font"] = {
                        "name": style.font.name,
                        "size": style.font.size.pt if style.font.size else None,
                        "bold": style.font.bold,
                        "italic": style.font.italic,
                        "underline": style.font.underline
                    }
                
                # Get paragraph format details if available
                if hasattr(style, "paragraph_format"):
                    pf = style.paragraph_format
                    info["paragraph_format"] = {
                        "alignment": str(pf.alignment).split('.')[-1] if pf.alignment else None,
                        "left_indent": pf.left_indent.inches if pf.left_indent else None,
                        "right_indent": pf.right_indent.inches if pf.right_indent else None,
                        "first_line_indent": pf.first_line_indent.inches if pf.first_line_indent else None,
                        "line_spacing": pf.line_spacing,
                        "space_before": pf.space_before.pt if pf.space_before else None,
                        "space_after": pf.space_after.pt if pf.space_after else None
                    }
                
                styles_info.append(info)
            
            return {
                "document": document_name,
                "styles": styles_info,
                "total": len(styles_info)
            }
        except FileNotFoundError as e:
            return {"error": str(e)}
        except Exception as e:
            return {"error": f"Error getting style details: {str(e)}"}
    
    @mcp.tool()
    def check_style_usage(document_name: str, style_name: str) -> Dict[str, Any]:
        """
        Checks where a style is used in a document.
        
        Args:
            document_name: Name of the document without extension
            style_name: Name of the style to check
            
        Returns:
            Dictionary with usage information
        """
        try:
            doc, _ = load_document(document_name)
            
            # Check if style exists
            try:
                _ = doc.styles[style_name]
            except KeyError:
                return {"error": f"Style '{style_name}' not found in document"}
            
            usage = {
                "style": style_name,
                "paragraph_usage": [],
                "table_usage": []
            }
            
            # Check paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.style.name == style_name:
                    usage["paragraph_usage"].append({
                        "index": i,
                        "text": para.text[:100] + ("..." if len(para.text) > 100 else "")
                    })
            
            # Check tables
            for t_idx, table in enumerate(doc.tables):
                for r_idx, row in enumerate(table.rows):
                    for c_idx, cell in enumerate(row.cells):
                        for p_idx, para in enumerate(cell.paragraphs):
                            if para.style.name == style_name:
                                usage["table_usage"].append({
                                    "table_index": t_idx,
                                    "row": r_idx,
                                    "col": c_idx,
                                    "paragraph": p_idx,
                                    "text": para.text[:50] + ("..." if len(para.text) > 50 else "")
                                })
            
            usage["total_usage"] = len(usage["paragraph_usage"]) + len(usage["table_usage"])
            return usage
        except FileNotFoundError as e:
            return {"error": str(e)}
        except Exception as e:
            return {"error": f"Error checking style usage: {str(e)}"}

    print("Document style tools registered")
    return mcp
