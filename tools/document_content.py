"""
Document Content Tools

This module provides tools for working with the content of Word documents,
including paragraphs, text, and basic content manipulation.
"""

from docx import Document
import os
from typing import Dict, List, Any, Optional

def register_document_content_tools(mcp):
    """
    Registers all document content-related tools with the MCP server.
    
    Args:
        mcp: The MCP server instance to register tools with
    """
    
    @mcp.tool()
    def add_paragraph(document_name: str, text: str, style: Optional[str] = None, 
                     formatting: Optional[Dict[str, Any]] = None) -> str:
        """
        Adds a paragraph to the document.
        
        Args:
            document_name: Name of the document without extension
            text: The text to add as a paragraph
            style: Optional style name to apply to the paragraph
            formatting: Optional dictionary of paragraph formatting properties
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            paragraph = doc.add_paragraph(text)
            
            if style:
                paragraph.style = style
                
            if formatting:
                # Apply paragraph formatting if provided
                if 'alignment' in formatting:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    alignment_map = {
                        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                        "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    if formatting['alignment'].upper() in alignment_map:
                        paragraph.alignment = alignment_map[formatting['alignment'].upper()]
            
            doc.save(doc_path)
            return f"Paragraph added to {document_name}.docx"
        except Exception as e:
            return f"Error adding paragraph: {str(e)}"
    
    @mcp.tool()
    def add_heading(document_name: str, heading_text: str, level: int = 1) -> str:
        """
        Adds a heading to the document.
        
        Args:
            document_name: Name of the document without extension
            heading_text: The heading text
            level: Heading level (0 for Title, 1 for Heading 1, etc.)
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            doc.add_heading(heading_text, level)
            doc.save(doc_path)
            return f"Heading added to {document_name}.docx"
        except Exception as e:
            return f"Error adding heading: {str(e)}"

    @mcp.tool()
    def get_paragraphs(document_name: str) -> List[Dict[str, Any]]:
        """
        Gets all paragraphs in the document.
        
        Args:
            document_name: Name of the document without extension
            
        Returns:
            List of paragraph details including text, style and index
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return [{"error": f"Document {document_name}.docx does not exist"}]
                
            doc = Document(doc_path)
            paragraphs = []
            
            for i, para in enumerate(doc.paragraphs):
                paragraphs.append({
                    "index": i,
                    "text": para.text,
                    "style": para.style.name
                })
            
            return paragraphs
        except Exception as e:
            return [{"error": f"Error getting paragraphs: {str(e)}"}]
    
    @mcp.tool()
    def get_paragraph_content(document_name: str, paragraph_index: int) -> Dict[str, Any]:
        """
        Gets the content of a specific paragraph.
        
        Args:
            document_name: Name of the document without extension
            paragraph_index: Index of the paragraph (0-based)
            
        Returns:
            Paragraph details including text and style
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return {"error": f"Document {document_name}.docx does not exist"}
                
            doc = Document(doc_path)
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return {"error": f"Paragraph index {paragraph_index} is out of range"}
                
            paragraph = doc.paragraphs[paragraph_index]
            return {
                "index": paragraph_index,
                "text": paragraph.text,
                "style": paragraph.style.name
            }
        except Exception as e:
            return {"error": f"Error getting paragraph: {str(e)}"}
    
    @mcp.tool()
    def add_formatted_text(document_name: str, paragraph_index: int, text: str, 
                          formatting: Dict[str, Any]) -> str:
        """
        Adds formatted text to an existing paragraph.
        
        Args:
            document_name: Name of the document without extension
            paragraph_index: Index of the paragraph to add text to (0-based)
            text: The text to add
            formatting: Dictionary of text formatting properties (bold, italic, etc.)
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return f"Error: Paragraph index {paragraph_index} is out of range"
                
            paragraph = doc.paragraphs[paragraph_index]
            run = paragraph.add_run(text)
            
            # Apply text formatting
            if 'bold' in formatting:
                run.bold = formatting['bold']
            if 'italic' in formatting:
                run.italic = formatting['italic']
            if 'underline' in formatting:
                run.underline = formatting['underline']
            if 'color' in formatting:
                run.font.color.rgb = formatting['color']
            if 'size' in formatting:
                run.font.size = formatting['size']
            if 'name' in formatting:
                run.font.name = formatting['name']
            
            doc.save(doc_path)
            return f"Formatted text added to paragraph {paragraph_index} in {document_name}.docx"
        except Exception as e:
            return f"Error adding formatted text: {str(e)}"
    
    @mcp.tool()
    def set_paragraph_text(document_name: str, paragraph_index: int, text: str) -> str:
        """
        Sets the text content of a specific paragraph.
        
        Args:
            document_name: Name of the document without extension
            paragraph_index: Index of the paragraph to modify (0-based)
            text: The new text content
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
                
            doc = Document(doc_path)
            if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
                return f"Error: Paragraph index {paragraph_index} is out of range"
                
            paragraph = doc.paragraphs[paragraph_index]
            
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ""
            
            # Add new text
            paragraph.add_run(text)
            
            doc.save(doc_path)
            return f"Text updated in paragraph {paragraph_index} in {document_name}.docx"
        except Exception as e:
            return f"Error setting paragraph text: {str(e)}"

    print("Document content tools registered")
    return mcp
