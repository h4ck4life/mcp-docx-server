from docx import Document
import os
from utils.helpers import get_document_path, load_document, _add_content_to_document

def register_document_core_tools(mcp):
    """Register all document core operation tools with the MCP server."""
    
    @mcp.resource("word://{doc_id}/content")
    def get_document_content(doc_id: str) -> str:
        """Reads the content of a Microsoft Word (.docx) document and returns it as text."""
        try:
            document = load_document(doc_id)
            full_text = [paragraph.text for paragraph in document.paragraphs]
            return '\n'.join(full_text)
        except ValueError as e:
            return str(e)
        except Exception as e:
            return f"Unexpected error: {str(e)}"

    @mcp.tool()
    def read_document(doc_id: str) -> str:
        """Reads the entire content of a Word document."""
        return get_document_content(doc_id)

    @mcp.tool()
    def check_document_exists(doc_id: str) -> str:
        """Checks if a Word document exists and can be read."""
        doc_path = get_document_path(doc_id)
        try:
            if os.path.exists(doc_path):
                document = Document(doc_path)
                paragraph_count = len(document.paragraphs)
                return f"Document '{doc_id}.docx' exists and is readable at path: {os.path.abspath(doc_path)}. Contains {paragraph_count} paragraphs."
            else:
                return f"Document '{doc_id}.docx' does not exist at path: {os.path.abspath(doc_path)}"
        except Exception as e:
            return f"Document '{doc_id}.docx' exists but cannot be read: {str(e)}"

    @mcp.tool()
    def list_available_documents() -> str:
        """Lists all Word documents (.docx files) available in the server directory."""
        try:
            script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
            docx_files = [f.replace('.docx', '') for f in os.listdir(script_dir) if f.endswith('.docx')]
            
            if not docx_files:
                return "No Word documents (.docx files) found in the server directory."
            
            doc_list = "\n".join([f"- {doc}" for doc in docx_files])
            return f"Available Word documents (without .docx extension):\n{doc_list}"
        except Exception as e:
            return f"Error listing documents: {str(e)}"
    
    @mcp.tool()
    def create_document(doc_id: str, title: str = "New Document") -> str:
        """Creates a new Word document with a title."""
        try:
            document = Document()
            document.add_heading(title, 0)
            doc_path = get_document_path(doc_id)
            document.save(doc_path)
            return f"Document '{doc_id}.docx' created successfully at path: {os.path.abspath(doc_path)}"
        except Exception as e:
            return f"Error creating document: {str(e)}"

    @mcp.tool()
    def update_document(doc_id: str, title: str = None, content: list = None, append: bool = True) -> str:
        """Updates an existing Word document by appending or replacing content."""
        try:
            doc_path = get_document_path(doc_id)
            
            if not os.path.exists(doc_path) and append:
                return f"Document '{doc_id}.docx' does not exist and cannot be updated. Create it first."
            
            if not append or not os.path.exists(doc_path):
                document = Document()
                if title:
                    document.add_heading(title, 0)
            else:
                document = Document(doc_path)
                if title:
                    document.add_heading(title, 1)
            
            if not _add_content_to_document(document, content):
                return "Error in table data: Number of data elements does not match table dimensions."
            
            document.save(doc_path)
            
            action = "updated by appending" if append else "replaced"
            title_msg = f" with new title" if title else ""
            content_msg = f" and {len(content) if content else 0} content items" if content else ""
            
            return f"Document '{doc_id}.docx' {action}{title_msg}{content_msg} successfully."
        except Exception as e:
            return f"Error updating document: {str(e)}"

    @mcp.tool()
    def append_to_document(doc_id: str, content: list) -> str:
        """Appends content to an existing Word document."""
        return update_document(doc_id, title=None, content=content, append=True)

    @mcp.tool()
    def replace_document(doc_id: str, title: str = None, content: list = None) -> str:
        """Replaces an existing Word document with new content."""
        return update_document(doc_id, title=title, content=content, append=False)
    
    print("Document Core tools registered.")
    return mcp
