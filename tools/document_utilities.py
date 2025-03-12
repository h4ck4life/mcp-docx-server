"""
Document Utility Tools

This module provides utility tools for Word documents such as conversion,
document analysis, and management functions.
"""

from docx import Document
import os
from typing import Dict, List, Any, Optional
import glob

def register_document_utility_tools(mcp):
    """
    Registers all document utility tools with the MCP server.
    
    Args:
        mcp: The MCP server instance to register tools with
    """
    
    @mcp.tool()
    def convert_to_pdf(document_name: str) -> str:
        """
        Converts a Word document to PDF format.
        
        Args:
            document_name: Name of the document without extension
            
        Returns:
            Message indicating success or failure
        """
        try:
            import comtypes.client
            
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
            
            pdf_path = f"{document_name}.pdf"
            
            # Create Word application
            word = comtypes.client.CreateObject('Word.Application')
            word.Visible = False
            
            # Open and convert document
            doc = word.Documents.Open(os.path.abspath(doc_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = PDF format
            doc.Close()
            word.Quit()
            
            return f"Document converted to PDF: {pdf_path}"
        except ImportError:
            return "Error: PDF conversion requires comtypes module. Please install it with 'pip install comtypes'"
        except Exception as e:
            return f"Error converting to PDF: {str(e)}"
    
    @mcp.tool()
    def analyze_document_structure(document_name: str) -> Dict[str, Any]:
        """
        Analyzes the structure of a document.
        
        Args:
            document_name: Name of the document without extension
            
        Returns:
            Document structure information
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return {"error": f"Document {document_name}.docx does not exist"}
            
            doc = Document(doc_path)
            
            # Analyze document structure
            structure = {
                "paragraphs": len(doc.paragraphs),
                "sections": len(doc.sections),
                "tables": len(doc.tables),
                "styles": [],
                "heading_structure": []
            }
            
            # Get styles info
            for style in doc.styles:
                if style.type == 1:  # 1 = paragraph style
                    structure["styles"].append(style.name)
            
            # Get heading structure
            for i, para in enumerate(doc.paragraphs):
                if "heading" in para.style.name.lower():
                    structure["heading_structure"].append({
                        "index": i,
                        "level": para.style.name.replace("Heading ", ""),
                        "text": para.text
                    })
            
            return structure
        except Exception as e:
            return {"error": f"Error analyzing document: {str(e)}"}
    
    @mcp.tool()
    def get_document_metadata(document_name: str) -> Dict[str, Any]:
        """
        Retrieves metadata information from a document.
        
        Args:
            document_name: Name of the document without extension
            
        Returns:
            Document metadata
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return {"error": f"Document {document_name}.docx does not exist"}
            
            doc = Document(doc_path)
            core_props = doc.core_properties
            
            metadata = {
                "author": core_props.author,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
                "title": core_props.title,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "category": core_props.category,
                "comments": core_props.comments,
                "last_modified_by": core_props.last_modified_by
            }
            
            return metadata
        except Exception as e:
            return {"error": f"Error getting metadata: {str(e)}"}
    
    @mcp.tool()
    def set_document_metadata(document_name: str, metadata: Dict[str, str]) -> str:
        """
        Sets metadata information for a document.
        
        Args:
            document_name: Name of the document without extension
            metadata: Dictionary with metadata fields to set
            
        Returns:
            Message indicating success
        """
        try:
            doc_path = f"{document_name}.docx"
            if not os.path.exists(doc_path):
                return f"Error: Document {document_name}.docx does not exist"
            
            doc = Document(doc_path)
            core_props = doc.core_properties
            
            # Set metadata fields
            if "author" in metadata:
                core_props.author = metadata["author"]
            if "title" in metadata:
                core_props.title = metadata["title"]
            if "subject" in metadata:
                core_props.subject = metadata["subject"]
            if "keywords" in metadata:
                core_props.keywords = metadata["keywords"]
            if "category" in metadata:
                core_props.category = metadata["category"]
            if "comments" in metadata:
                core_props.comments = metadata["comments"]
            
            doc.save(doc_path)
            return f"Metadata updated for {document_name}.docx"
        except Exception as e:
            return f"Error setting metadata: {str(e)}"
    
    print("Document utility tools registered")
    return mcp
