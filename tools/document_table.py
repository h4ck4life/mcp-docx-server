from docx.table import _Cell, _Row, Table
from docx.shared import Pt, Inches
from utils.helpers import load_document, save_document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def parse_color(color_str):
    """Parse color string (hex or rgb) into RGB components."""
    if color_str.startswith('#'):
        # Convert hex color to RGB
        r = int(color_str[1:3], 16)
        g = int(color_str[3:5], 16)
        b = int(color_str[5:7], 16)
        return (r, g, b)
    elif color_str.startswith('rgb('):
        # Parse rgb() format
        rgb = color_str.strip('rgb()').split(',')
        if len(rgb) == 3:
            r = int(rgb[0].strip())
            g = int(rgb[1].strip())
            b = int(rgb[2].strip())
            return (r, g, b)
    return (0, 0, 0)  # Default to black

def register_document_table_tools(mcp):
    """Register all table-related tools with the MCP server."""
    
    # Initialize registered_tools set if it doesn't exist
    if not hasattr(mcp, 'registered_tools'):
        mcp.registered_tools = set()
    
    # Check if tools are already registered to avoid duplicates
    registered_tools = mcp.registered_tools
    
    if 'add_table' not in registered_tools:
        @mcp.tool()
        def add_table(document_name: str, rows: int, cols: int, data: str = None, style: str = None) -> int:
            """
            Adds a new table to the document.
            
            Args:
                document_name: The name of the document (without .docx extension)
                rows: Number of rows
                cols: Number of columns
                data: Optional comma-separated values to fill the table with (row by row)
                style: Optional table style name
                
            Returns:
                The index of the newly added table or -1 if failed
            """
            doc = load_document(document_name)
            
            try:
                # Add table to the document
                table = doc.add_table(rows=rows, cols=cols)
                
                # Apply style if provided
                if style:
                    try:
                        table.style = style
                    except:
                        # Style doesn't exist, continue without styling
                        pass
                
                # Fill with data if provided
                if data:
                    values = data.split(',')
                    cell_idx = 0
                    
                    for i in range(rows):
                        for j in range(cols):
                            if cell_idx < len(values):
                                table.cell(i, j).text = values[cell_idx].strip()
                            cell_idx += 1
                
                # Get the table index
                table_idx = len(doc.tables) - 1
                
                # Save the document
                save_document(doc, document_name)
                return table_idx
            except Exception as e:
                print(f"Error adding table: {e}")
                return -1
        registered_tools.add('add_table')
    
    if 'list_tables' not in registered_tools:
        @mcp.tool()
        def list_tables(document_name: str) -> list:
            """
            Lists all tables in the document with basic information.
            
            Args:
                document_name: The name of the document (without .docx extension)
                
            Returns:
                List of dictionaries with table information
            """
            doc = load_document(document_name)
            tables_info = []
            
            for i, table in enumerate(doc.tables):
                tables_info.append({
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "style": table.style.name if table.style else "None"
                })
                
            return tables_info
        registered_tools.add('list_tables')
    
    if 'get_table_data' not in registered_tools:
        @mcp.tool()
        def get_table_data(document_name: str, table_index: int) -> list:
            """
            Gets the content of a table as a 2D array.
            
            Args:
                document_name: The name of the document (without .docx extension)
                table_index: The index of the table to retrieve
                
            Returns:
                2D list representing the table data
            """
            doc = load_document(document_name)
            
            try:
                if table_index < 0 or table_index >= len(doc.tables):
                    return []
                    
                table = doc.tables[table_index]
                data = []
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    data.append(row_data)
                    
                return data
            except Exception as e:
                print(f"Error retrieving table data: {e}")
                return []
        registered_tools.add('get_table_data')
    
    if 'merge_table_cells' not in registered_tools:
        @mcp.tool()
        def merge_table_cells(document_name: str, table_index: int,
                             start_row: int, start_col: int, 
                             end_row: int, end_col: int) -> bool:
            """
            Merges cells in a table.
            
            Args:
                document_name: The name of the document (without .docx extension)
                table_index: The index of the table
                start_row: Starting row index (0-based)
                start_col: Starting column index (0-based)
                end_row: Ending row index (0-based)
                end_col: Ending column index (0-based)
                
            Returns:
                True if successful, False otherwise
            """
            doc = load_document(document_name)
            
            try:
                if table_index < 0 or table_index >= len(doc.tables):
                    return False
                    
                table = doc.tables[table_index]
                 
                # Check bounds
                if (start_row < 0 or start_row >= len(table.rows) or
                    end_row < 0 or end_row >= len(table.rows) or
                    start_col < 0 or start_col >= len(table.columns) or
                    end_col < 0 or end_col >= len(table.columns)):
                    return False
                    
                # Get the first cell
                start_cell = table.cell(start_row, start_col)
                
                # Merge cells
                start_cell.merge(table.cell(end_row, end_col))
                
                save_document(doc, document_name)
                return True
            except Exception as e:
                print(f"Error merging cells: {e}")
                return False
        registered_tools.add('merge_table_cells')
    
    if 'set_table_column_width' not in registered_tools:
        @mcp.tool()
        def set_table_column_width(document_name: str, table_index: int,
                                  column_index: int, width: float) -> bool:
            """
            Sets the width of a table column in inches.
            
            Args:
                document_name: The name of the document (without .docx extension)
                table_index: The index of the table
                column_index: The index of the column to modify
                width: Width in inches
                
            Returns:
                True if successful, False otherwise
            """
            doc = load_document(document_name)
            
            try:
                if table_index < 0 or table_index >= len(doc.tables):
                    return False
                    
                table = doc.tables[table_index]
                
                # Check bounds
                if column_index < 0 or column_index >= len(table.columns):
                    return False
                    
                # Set width for each cell in the column
                for cell in table.columns[column_index].cells:
                    cell.width = Inches(width)
                
                save_document(doc, document_name)
                return True
            except Exception as e:
                print(f"Error setting column width: {e}")
                return False
        registered_tools.add('set_table_column_width')
    
    if 'set_table_cell_properties' not in registered_tools:
        @mcp.tool()
        def set_table_cell_properties(document_name: str, table_index: int,
                                     row_index: int, col_index: int,
                                     properties: dict) -> bool:
            """
            Sets properties for a specific table cell.
            
            Args:
                document_name: The name of the document (without .docx extension)
                table_index: The index of the table
                row_index: The row index of the cell
                col_index: The column index of the cell
                properties: Dictionary of properties to set:
                    {
                        "text": "New cell content",
                        "background_color": "#FFFF00",
                        "vertical_alignment": "CENTER"  # TOP, CENTER, BOTTOM
                    }
                
            Returns:
                True if successful, False otherwise
            """
            doc = load_document(document_name)
            
            try:
                if table_index < 0 or table_index >= len(doc.tables):
                    return False
                    
                table = doc.tables[table_index]
                
                # Check bounds
                if (row_index < 0 or row_index >= len(table.rows) or
                    col_index < 0 or col_index >= len(table.columns)):
                    return False
                    
                cell = table.cell(row_index, col_index)
                
                # Set text content if provided
                if "text" in properties:
                    cell.text = properties["text"]
                
                # Set background color if provided
                if "background_color" in properties:
                    rgb = parse_color(properties["background_color"])
                    cell_properties = cell._element.tcPr
                    shading_element = cell_properties.xpath('./w:shd')
                    
                    if not shading_element:
                        shading_element = OxmlElement('w:shd')
                        cell_properties.append(shading_element)
                    
                    shading_element.set(qn('w:fill'), f"{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}")
                
                # Set vertical alignment if provided
                if "vertical_alignment" in properties:
                    alignment_map = {
                        "TOP": 0,
                        "CENTER": 1,
                        "BOTTOM": 2
                    }
                    cell.vertical_alignment = alignment_map.get(properties["vertical_alignment"].upper(), 0)
                
                save_document(doc, document_name)
                return True
            except Exception as e:
                print(f"Error setting cell properties: {e}")
                return False
        registered_tools.add('set_table_cell_properties')
    
    # No need for the update call at the end since we're adding to registered_tools directly
    print("Document Table tools registered.")
    return mcp
