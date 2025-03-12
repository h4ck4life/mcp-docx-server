from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import os

# Helper functions for document operations
def get_document_path(doc_id: str) -> str:
    """Returns the full path to a document in the same directory as this script."""
    script_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(script_dir, f"{doc_id}.docx")

def load_document(doc_id: str) -> Document:
    """Loads a Word document, handling potential FileNotFoundError."""
    doc_path = get_document_path(doc_id)
    try:
        return Document(doc_path)
    except FileNotFoundError:
        raise ValueError(f"Document '{doc_id}.docx' not found.")
    except Exception as e:
        raise ValueError(f"Error loading document '{doc_id}.docx': {str(e)}")

def style_exists(document, style_name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    """Checks if a style exists in the document."""
    for style in document.styles:
        if style.name == style_name and style.type == style_type:
            return True
    return False

# Formatting helpers
def _apply_paragraph_formatting(paragraph, formatting):
    """Apply formatting to a paragraph."""
    if not formatting:
        return
    
    para_format = paragraph.paragraph_format
    
    # Alignment
    alignment = formatting.get("alignment")
    if alignment:
        alignment_map = {
            "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
            "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
            "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
            "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if alignment.upper() in alignment_map:
            para_format.alignment = alignment_map[alignment.upper()]
    
    # Indentation
    left_indent = formatting.get("left_indent")
    if left_indent is not None:
        para_format.left_indent = Inches(float(left_indent))
    
    right_indent = formatting.get("right_indent")
    if right_indent is not None:
        para_format.right_indent = Inches(float(right_indent))
    
    first_line_indent = formatting.get("first_line_indent")
    if first_line_indent is not None:
        para_format.first_line_indent = Inches(float(first_line_indent))
    
    # Spacing
    space_before = formatting.get("space_before")
    if space_before is not None:
        para_format.space_before = Pt(float(space_before))
    
    space_after = formatting.get("space_after")
    if space_after is not None:
        para_format.space_after = Pt(float(space_after))
    
    # Line spacing
    line_spacing = formatting.get("line_spacing")
    if line_spacing is not None:
        try:
            # Try as a float (multiple)
            spacing_float = float(line_spacing)
            para_format.line_spacing = spacing_float
        except ValueError:
            # Try as a point value
            para_format.line_spacing = Pt(float(line_spacing))
    
    # Pagination
    keep_together = formatting.get("keep_together")
    if keep_together is not None:
        para_format.keep_together = bool(keep_together)
    
    keep_with_next = formatting.get("keep_with_next")
    if keep_with_next is not None:
        para_format.keep_with_next = bool(keep_with_next)
    
    page_break_before = formatting.get("page_break_before")
    if page_break_before is not None:
        para_format.page_break_before = bool(page_break_before)
    
    widow_control = formatting.get("widow_control")
    if widow_control is not None:
        para_format.widow_control = bool(widow_control)

def _apply_run_formatting(run, formatting):
    """Apply formatting to a run of text."""
    if not formatting:
        return
    
    font = run.font
    
    # Font name and size
    font_name = formatting.get("name")
    if font_name:
        font.name = font_name
    
    font_size = formatting.get("size")
    if font_size is not None:
        font.size = Pt(float(font_size))
    
    # Font styles
    bold = formatting.get("bold")
    if bold is not None:
        font.bold = bool(bold)
    
    italic = formatting.get("italic")
    if italic is not None:
        font.italic = bool(italic)
    
    underline = formatting.get("underline")
    if underline is not None:
        font.underline = bool(underline)
    
    # Font color
    color = formatting.get("color")
    if color:
        if color.startswith('#'):
            # Convert hex color to RGB
            r = int(color[1:3], 16)
            g = int(color[3:5], 16)
            b = int(color[5:7], 16)
            font.color.rgb = RGBColor(r, g, b)
        elif color.startswith('rgb('):
            # Parse rgb() format
            rgb = color.strip('rgb()').split(',')
            if len(rgb) == 3:
                r = int(rgb[0].strip())
                g = int(rgb[1].strip())
                b = int(rgb[2].strip())
                font.color.rgb = RGBColor(r, g, b)

def _add_content_to_document(document, content):
    """Helper function to add content to a document object."""
    if not content:
        return True
        
    for item in content:
        content_type = item.get("type", "").lower()
        text = item.get("text", "")
        
        if content_type == "heading":
            level = item.get("level", 1)
            heading = document.add_heading(text, level)
            
            # Apply formatting if provided
            formatting = item.get("formatting", {})
            if formatting:
                _apply_paragraph_formatting(heading, formatting)
            
        elif content_type == "paragraph":
            style = item.get("style")
            try:
                # If style is specified, ensure it exists in the document
                if style:
                    style_exists_in_doc = False
                    for doc_style in document.styles:
                        if doc_style.name == style:
                            style_exists_in_doc = True
                            break
                    
                    # If style doesn't exist, it might be a built-in style that needs defining
                    if not style_exists_in_doc:
                        # Add temporary paragraph to define the style
                        temp_para = document.add_paragraph("", style=style)
                        # Remove the temporary paragraph
                        p = temp_para._element
                        p.getparent().remove(p)
                
                paragraph = document.add_paragraph(text, style=style if style else None)
                
                # Apply formatting if provided
                formatting = item.get("formatting", {})
                if formatting:
                    _apply_paragraph_formatting(paragraph, formatting)
                
                # Apply run formatting if provided
                run_formatting = item.get("run_formatting", {})
                if run_formatting and len(paragraph.runs) > 0:
                    _apply_run_formatting(paragraph.runs[0], run_formatting)
                
            except KeyError:
                # Style not found, add without style
                paragraph = document.add_paragraph(text)
                
        elif content_type == "table":
            rows = item.get("rows", 1)
            cols = item.get("cols", 1)
            data = item.get("data", "")
            style = item.get("style")
            
            table = document.add_table(rows=rows, cols=cols)
            
            # Apply style if specified
            if style:
                try:
                    # If style doesn't exist, it might be a built-in style that needs defining
                    style_exists_in_doc = False
                    for doc_style in document.styles:
                        if doc_style.name == style and doc_style.type == WD_STYLE_TYPE.TABLE:
                            style_exists_in_doc = True
                            break
                    
                    # If style doesn't exist, it might be a built-in style that needs defining
                    if not style_exists_in_doc:
                        # Add temporary table to define the style
                        temp_table = document.add_table(rows=1, cols=1)
                        temp_table.style = style
                        # Remove the temporary table
                        p = temp_table._element
                        p.getparent().remove(p)
                    
                    table.style = style
                except KeyError:
                    pass  # Continue without style if not found
            
            # Fill with data if provided
            if data:
                data_list = data.split(',')
                
                # Pad with empty strings if too few data elements
                if len(data_list) < rows * cols:
                    data_list.extend([''] * (rows * cols - len(data_list)))
                    
                # Check if data matches table dimensions
                if len(data_list) != rows * cols:
                    return False
                
                # Fill table cells
                for i in range(rows):
                    for j in range(cols):
                        cell_idx = i * cols + j
                        if cell_idx < len(data_list):
                            table.cell(i, j).text = data_list[cell_idx].strip()
            
            # Process cell_formatting if provided
            cell_formatting = item.get("cell_formatting", [])
            for cell_format in cell_formatting:
                row = cell_format.get("row", 0)
                col = cell_format.get("col", 0)
                formatting = cell_format.get("formatting", {})
                
                if row < rows and col < cols:
                    cell = table.cell(row, col)
                    if cell and len(cell.paragraphs) > 0:
                        _apply_paragraph_formatting(cell.paragraphs[0], formatting)
    
    return True

def save_document(document, document_name):
    """
    Save a document to the documents folder.
    
    Args:
        document: The document object to save
        document_name: The name of the document (without .docx extension)
    """
    import os
    from pathlib import Path
    
    # Ensure document_name doesn't have .docx extension
    if document_name.lower().endswith('.docx'):
        document_name = document_name[:-5]
    
    # Determine the documents directory
    documents_dir = os.environ.get('DOCX_DOCUMENTS_DIR', 'documents')
    Path(documents_dir).mkdir(exist_ok=True)
    
    # Save the document
    document.save(os.path.join(documents_dir, f"{document_name}.docx"))
