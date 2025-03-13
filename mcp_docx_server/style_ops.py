"""
Style management operations for Word documents.
"""

from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mcp_docx_server.utils import load_document, get_document_path, style_exists

def ensure_style_exists(doc_id: str, style_name: str, style_type: str = "paragraph") -> str:
    """
    Ensures a style exists in the document by applying it to a temporary paragraph.
    
    This is useful for built-in styles that need to be defined in the document before use.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        style_name (str): The name of the style to ensure exists.
        style_type (str): Type of style: 'paragraph', 'character', or 'table'.
    
    Returns:
        str: Status message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        # Map string type to enum
        style_type_map = {
            "paragraph": WD_STYLE_TYPE.PARAGRAPH,
            "character": WD_STYLE_TYPE.CHARACTER,
            "table": WD_STYLE_TYPE.TABLE
        }
        
        if style_type.lower() not in style_type_map:
            return f"Error: Invalid style type '{style_type}'. Valid values are: {', '.join(style_type_map.keys())}"
        
        style_type_enum = style_type_map[style_type.lower()]
        
        # Check if style already exists
        if style_exists(document, style_name, style_type_enum):
            return f"Style '{style_name}' already exists in document."
        
        # For paragraph style, add a temporary paragraph with the style
        if style_type_enum == WD_STYLE_TYPE.PARAGRAPH:
            try:
                temp_para = document.add_paragraph("Style Definition", style=style_name)
                # Remove the temporary paragraph
                p = temp_para._element
                p.getparent().remove(p)
                
                doc_path = get_document_path(doc_id)
                document.save(doc_path)
                return f"Paragraph style '{style_name}' successfully defined in document."
            except KeyError:
                return f"Error: Built-in style '{style_name}' not found in Word. Check the style name."
        
        # For character style, add a paragraph with a run using the style
        elif style_type_enum == WD_STYLE_TYPE.CHARACTER:
            try:
                temp_para = document.add_paragraph()
                temp_run = temp_para.add_run("Style Definition")
                temp_run.style = style_name
                
                # Remove the temporary paragraph
                p = temp_para._element
                p.getparent().remove(p)
                
                doc_path = get_document_path(doc_id)
                document.save(doc_path)
                return f"Character style '{style_name}' successfully defined in document."
            except KeyError:
                return f"Error: Built-in style '{style_name}' not found in Word. Check the style name."
        
        # For table style, add a temporary table with the style
        elif style_type_enum == WD_STYLE_TYPE.TABLE:
            try:
                temp_table = document.add_table(rows=1, cols=1)
                temp_table.style = style_name
                
                # Remove the temporary table
                p = temp_table._element
                p.getparent().remove(p)
                
                doc_path = get_document_path(doc_id)
                document.save(doc_path)
                return f"Table style '{style_name}' successfully defined in document."
            except KeyError:
                return f"Error: Built-in style '{style_name}' not found in Word. Check the style name."
        
        return "Unknown error ensuring style exists."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error ensuring style exists: {str(e)}"

def create_custom_style(doc_id: str, style_name: str, style_type: str = "paragraph", base_style: str = None) -> str:
    """
    Creates a custom style in the document.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        style_name (str): The name for the new style.
        style_type (str): Type of style: 'paragraph', 'character', or 'table'.
        base_style (str, optional): Name of the style to base this one on.
    
    Returns:
        str: Status message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        # Map string type to enum
        style_type_map = {
            "paragraph": WD_STYLE_TYPE.PARAGRAPH,
            "character": WD_STYLE_TYPE.CHARACTER,
            "table": WD_STYLE_TYPE.TABLE
        }
        
        if style_type.lower() not in style_type_map:
            return f"Error: Invalid style type '{style_type}'. Valid values are: {', '.join(style_type_map.keys())}"
        
        style_type_enum = style_type_map[style_type.lower()]
        
        # Check if style already exists
        if style_exists(document, style_name, style_type_enum):
            return f"Error: Style '{style_name}' already exists in document."
        
        # If base_style is provided, ensure it exists first
        if base_style:
            if not style_exists(document, base_style, style_type_enum):
                # Try to ensure base style exists if it's a built-in style
                ensure_result = ensure_style_exists(doc_id, base_style, style_type)
                if "Error" in ensure_result or "not found" in ensure_result:
                    return f"Error: Base style '{base_style}' does not exist and could not be defined."
        
        # Create the new style
        new_style = document.styles.add_style(style_name, style_type_enum)
        
        # Set base style if provided
        if base_style:
            try:
                new_style.base_style = document.styles[base_style]
            except KeyError:
                return f"Error setting base style: Style '{base_style}' not found."
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Custom {style_type} style '{style_name}' created successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error creating custom style: {str(e)}"

def modify_style(doc_id: str, style_name: str, properties: dict) -> str:
    """
    Modifies properties of an existing style.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        style_name (str): The name of the style to modify.
        properties (dict): Dictionary with style properties to modify:
            - font: Font properties dict (name, size, bold, italic, underline, color)
            - paragraph: Paragraph formatting properties dict
    
    Returns:
        str: Status message indicating success or failure.
    """
    try:
        document = load_document(doc_id)
        
        # Check if style exists
        try:
            style = document.styles[style_name]
        except KeyError:
            return f"Error: Style '{style_name}' not found in document."
        
        # Modify font properties if provided
        font_props = properties.get("font", {})
        if font_props and hasattr(style, "font"):
            font = style.font
            
            # Font name
            if "name" in font_props:
                font.name = font_props["name"]
            
            # Font size
            if "size" in font_props:
                font.size = Pt(float(font_props["size"]))
            
            # Font styles
            if "bold" in font_props:
                font.bold = bool(font_props["bold"])
            
            if "italic" in font_props:
                font.italic = bool(font_props["italic"])
            
            if "underline" in font_props:
                font.underline = bool(font_props["underline"])
            
            # Font color
            if "color" in font_props:
                color = font_props["color"]
                if color.startswith('#'):
                    # Convert hex color to RGB
                    r = int(color[1:3], 16)
                    g = int(color[3:5], 16)
                    b = int(color[5:7], 16)
                    font.color.rgb = RGBColor(r, g, b)
                elif color.startswith('rgb('):
                    # Parse rgb() format
                    rgb = color.strip('rgb()').split(',')
                    if len(rgb) == 3:
                        r = int(rgb[0].strip())
                        g = int(rgb[1].strip())
                        b = int(rgb[2].strip())
                        font.color.rgb = RGBColor(r, g, b)
        
        # Modify paragraph formatting properties if provided
        para_props = properties.get("paragraph", {})
        if para_props and hasattr(style, "paragraph_format"):
            para_format = style.paragraph_format
            
            # Alignment
            alignment = para_props.get("alignment")
            if alignment:
                alignment_map = {
                    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
                    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY
                }
                if alignment.upper() in alignment_map:
                    para_format.alignment = alignment_map[alignment.upper()]
            
            # Indentation
            left_indent = para_props.get("left_indent")
            if left_indent is not None:
                para_format.left_indent = Inches(float(left_indent))
            
            right_indent = para_props.get("right_indent")
            if right_indent is not None:
                para_format.right_indent = Inches(float(right_indent))
            
            first_line_indent = para_props.get("first_line_indent")
            if first_line_indent is not None:
                para_format.first_line_indent = Inches(float(first_line_indent))
            
            # Spacing
            space_before = para_props.get("space_before")
            if space_before is not None:
                para_format.space_before = Pt(float(space_before))
            
            space_after = para_props.get("space_after")
            if space_after is not None:
                para_format.space_after = Pt(float(space_after))
            
            # Line spacing
            line_spacing = para_props.get("line_spacing")
            if line_spacing is not None:
                try:
                    # Try as a float (multiple)
                    spacing_float = float(line_spacing)
                    para_format.line_spacing = spacing_float
                except ValueError:
                    # Try as a point value
                    para_format.line_spacing = Pt(float(line_spacing))
            
            # Pagination
            keep_together = para_props.get("keep_together")
            if keep_together is not None:
                para_format.keep_together = bool(keep_together)
            
            keep_with_next = para_props.get("keep_with_next")
            if keep_with_next is not None:
                para_format.keep_with_next = bool(keep_with_next)
            
            page_break_before = para_props.get("page_break_before")
            if page_break_before is not None:
                para_format.page_break_before = bool(page_break_before)
            
            widow_control = para_props.get("widow_control")
            if widow_control is not None:
                para_format.widow_control = bool(widow_control)
        
        # Additional style properties
        if "quick_style" in properties:
            style.quick_style = bool(properties["quick_style"])
        
        if "hidden" in properties:
            style.hidden = bool(properties["hidden"])
        
        if "priority" in properties:
            style.priority = int(properties["priority"])
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Style '{style_name}' modified successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error modifying style: {str(e)}"

def get_styles_detail(doc_id: str, style_type: str = None) -> str:
    """
    Gets detailed information about styles in the document.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        style_type (str, optional): Filter by style type: 'paragraph', 'character', or 'table'.
                                   If None, shows all styles.
    
    Returns:
        str: Detailed information about styles in the document.
    """
    try:
        document = load_document(doc_id)
        
        # Map string type to enum if provided
        style_type_enum = None
        if style_type:
            style_type_map = {
                "paragraph": WD_STYLE_TYPE.PARAGRAPH,
                "character": WD_STYLE_TYPE.CHARACTER,
                "table": WD_STYLE_TYPE.TABLE
            }
            
            if style_type.lower() not in style_type_map:
                return f"Error: Invalid style type '{style_type}'. Valid values are: {', '.join(style_type_map.keys())}"
            
            style_type_enum = style_type_map[style_type.lower()]
        
        # Get styles
        styles_info = []
        for style in document.styles:
            # Skip if filtering by type and this style doesn't match
            if style_type_enum and style.type != style_type_enum:
                continue
            
            # Get style type as string
            style_type_str = "Unknown"
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                style_type_str = "Paragraph"
            elif style.type == WD_STYLE_TYPE.CHARACTER:
                style_type_str = "Character"
            elif style.type == WD_STYLE_TYPE.TABLE:
                style_type_str = "Table"
            elif style.type == WD_STYLE_TYPE.LIST:
                style_type_str = "List"
            
            # Get base style name if available
            base_style = "None"
            if hasattr(style, "base_style") and style.base_style:
                base_style = style.base_style.name
            
            # Create style info
            style_info = [
                f"Style: {style.name}",
                f"  Type: {style_type_str}",
                f"  Base Style: {base_style}"
            ]
            
            # Add behavior properties if available
            behavior_props = []
            if hasattr(style, "quick_style"):
                behavior_props.append(f"Quick Style: {style.quick_style}")
            if hasattr(style, "priority"):
                behavior_props.append(f"Priority: {style.priority}")
            if hasattr(style, "hidden"):
                behavior_props.append(f"Hidden: {style.hidden}")
            
            if behavior_props:
                style_info.append("  Behavior:")
                for prop in behavior_props:
                    style_info.append(f"    {prop}")
            
            styles_info.append("\n".join(style_info))
        
        if not styles_info:
            return f"No styles found in document{' with type ' + style_type if style_type else ''}."
        
        return "\n\n".join(styles_info)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error getting styles detail: {str(e)}"

def check_style_usage(doc_id: str, style_name: str) -> str:
    """
    Checks if a style is used in the document and where.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        style_name (str): The name of the style to check.
    
    Returns:
        str: Information about where the style is used.
    """
    try:
        document = load_document(doc_id)
        
        # Check if style exists
        try:
            style = document.styles[style_name]
        except KeyError:
            return f"Style '{style_name}' not found in document."
        
        # Get style type
        style_type_str = "Unknown"
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style_type_str = "Paragraph"
        elif style.type == WD_STYLE_TYPE.CHARACTER:
            style_type_str = "Character"
        elif style.type == WD_STYLE_TYPE.TABLE:
            style_type_str = "Table"
        
        # Check usage based on style type
        usage_locations = []
        
        # Check paragraphs for paragraph and character styles
        if style.type in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
            for i, para in enumerate(document.paragraphs):
                # Check paragraph style
                if style.type == WD_STYLE_TYPE.PARAGRAPH and para.style and para.style.name == style_name:
                    preview = para.text[:30] + ("..." if len(para.text) > 30 else "")
                    usage_locations.append(f"Paragraph {i}: \"{preview}\"")
                
                # Check character styles in runs
                if style.type == WD_STYLE_TYPE.CHARACTER:
                    for j, run in enumerate(para.runs):
                        if run.style and run.style.name == style_name:
                            preview = run.text[:30] + ("..." if len(run.text) > 30 else "")
                            usage_locations.append(f"Paragraph {i}, Run {j}: \"{preview}\"")
        
        # Check tables for table styles
        if style.type == WD_STYLE_TYPE.TABLE:
            for i, table in enumerate(document.tables):
                if table.style and table.style.name == style_name:
                    rows = len(table.rows)
                    cols = len(table.rows[0].cells) if rows > 0 else 0
                    usage_locations.append(f"Table {i}: {rows}x{cols} table")
        
        # Report results
        if not usage_locations:
            return f"{style_type_str} style '{style_name}' exists in the document but is not currently used."
        
        return f"{style_type_str} style '{style_name}' is used in the following locations:\n" + "\n".join(usage_locations)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error checking style usage: {str(e)}"

def list_styles(doc_id: str) -> str:
    """Lists available paragraph and character styles in the document."""
    try:
        document = load_document(doc_id)
        para_styles = []
        char_styles = []
        table_styles = []
        
        for style in document.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                para_styles.append(style.name)
            elif style.type == WD_STYLE_TYPE.CHARACTER:
                char_styles.append(style.name)
            elif style.type == WD_STYLE_TYPE.TABLE:
                table_styles.append(style.name)
        
        result = []
        if para_styles:
            result.append("Paragraph styles:\n" + ", ".join(para_styles))
        if char_styles:
            result.append("\nCharacter styles:\n" + ", ".join(char_styles))
        if table_styles:
            result.append("\nTable styles:\n" + ", ".join(table_styles))
            
        return "\n".join(result)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error listing styles: {str(e)}"
