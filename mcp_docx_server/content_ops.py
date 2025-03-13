"""
Content operations for Word documents (paragraphs, tables, etc.)
"""

from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
import base64
from io import BytesIO
from mcp_docx_server.utils import load_document, get_document_path, apply_paragraph_formatting, apply_run_formatting

def add_paragraph(doc_id: str, text: str, style: str = None, formatting: dict = None) -> str:
    """Adds a paragraph to an existing Word document, optionally with style and formatting.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        text (str): The paragraph text.
        style (str, optional): The paragraph style.
        formatting (dict, optional): Dictionary with paragraph formatting options like:
            - alignment: 'LEFT', 'CENTER', 'RIGHT', 'JUSTIFY'
            - left_indent, right_indent, first_line_indent: Indentation in inches
            - space_before, space_after: Spacing in points
            - line_spacing: Line spacing as multiple or points
            - keep_together, keep_with_next, page_break_before, widow_control: Boolean pagination options
    """
    try:
        document = load_document(doc_id)
        result_message = "Paragraph added successfully."
        
        # If style is specified, ensure it exists in the document
        if style:
            style_exists_in_doc = False
            for doc_style in document.styles:
                if doc_style.name == style and doc_style.type == WD_STYLE_TYPE.PARAGRAPH:
                    style_exists_in_doc = True
                    break
            
            # If style doesn't exist, it might be a built-in style that needs defining
            if not style_exists_in_doc:
                try:
                    # Add temporary paragraph to define the style
                    temp_para = document.add_paragraph("", style=style)
                    # Remove the temporary paragraph
                    p = temp_para._element
                    p.getparent().remove(p)
                except KeyError:
                    result_message = f"Warning: Style '{style}' not found. Added without style."
                    style = None
        
        # Now add the actual paragraph
        if style:
            try:
                paragraph = document.add_paragraph(text, style=style)
            except KeyError:
                paragraph = document.add_paragraph(text)
                result_message = f"Warning: Style '{style}' not found. Added without style."
        else:
            paragraph = document.add_paragraph(text)
        
        # Apply formatting if provided
        if formatting:
            apply_paragraph_formatting(paragraph, formatting)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return result_message
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding paragraph: {str(e)}"

def add_formatted_text(doc_id: str, paragraph_index: int, text: str, formatting: dict = None) -> str:
    """Adds formatted text to an existing paragraph.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        paragraph_index (int): The index of the paragraph to add text to (0-based).
        text (str): The text to add.
        formatting (dict, optional): Dictionary with text formatting options like:
            - name: Font name
            - size: Font size in points
            - bold, italic, underline: Boolean style options
            - color: Color as hex (#RRGGBB) or rgb(r,g,b)
    """
    try:
        document = load_document(doc_id)
        
        if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
            return "Error: Paragraph index out of range."
        
        paragraph = document.paragraphs[paragraph_index]
        run = paragraph.add_run(text)
        
        # Apply formatting if provided
        if formatting:
            apply_run_formatting(run, formatting)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return "Formatted text added successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding formatted text: {str(e)}"

def add_image(doc_id: str, image_data: str, image_name: str, width_inches: float = 6.0) -> str:
    """Adds an image to an existing Word document."""
    try:
        document = load_document(doc_id)
        image_bytes = base64.b64decode(image_data)
        image_stream = BytesIO(image_bytes)
        document.add_picture(image_stream, width=Inches(width_inches))
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Image '{image_name}' added to document '{doc_id}.docx' successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding image: {str(e)}"

def add_heading(doc_id: str, text: str, level: int, formatting: dict = None) -> str:
    """Adds a heading to an existing Word document with optional formatting.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        text (str): The heading text.
        level (int): The heading level (0-9, where 0 is Title).
        formatting (dict, optional): Dictionary with paragraph formatting options.
    """
    try:
        document = load_document(doc_id)
        
        # First, ensure the heading style exists in the document
        heading_style = None
        if level == 0:
            heading_style = "Title"
        else:
            heading_style = f"Heading {level}"
        
        # Check if this heading style exists in the document
        style_exists_in_doc = False
        for doc_style in document.styles:
            if doc_style.name == heading_style:
                style_exists_in_doc = True
                break
        
        # If style doesn't exist, it needs to be defined first
        if not style_exists_in_doc:
            try:
                # Add temporary paragraph to define the style
                temp_para = document.add_paragraph("", style=heading_style)
                # Remove the temporary paragraph
                p = temp_para._element
                p.getparent().remove(p)
            except KeyError:
                # If the style is not found, it's not a built-in style
                pass
        
        # Now add the actual heading
        heading = document.add_heading(text, level)
        
        # Apply formatting if provided
        if formatting:
            apply_paragraph_formatting(heading, formatting)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return "Heading added successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding heading: {str(e)}"

def add_table(doc_id: str, rows: int, cols: int, data: str = None, style: str = None) -> str:
    """Adds a table to an existing Word document.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        rows (int): The number of rows in the table.
        cols (int): The number of columns in the table.
        data (str, optional): Comma-separated data for the table cells (row-wise).
        style (str, optional): Table style (e.g., "Table Grid", "Light Shading").
    """
    try:
        document = load_document(doc_id)
        
        # If style is specified, ensure it exists in the document
        if style:
            style_exists_in_doc = False
            for doc_style in document.styles:
                if doc_style.name == style and doc_style.type == WD_STYLE_TYPE.TABLE:
                    style_exists_in_doc = True
                    break
            
            # If style doesn't exist, it might be a built-in style that needs defining
            if not style_exists_in_doc:
                try:
                    # Add temporary table to define the style
                    temp_table = document.add_table(rows=1, cols=1)
                    temp_table.style = style
                    # Remove the temporary table
                    p = temp_table._element
                    p.getparent().remove(p)
                except KeyError:
                    return f"Warning: Table style '{style}' not found. Table will be added with default style."
        
        # Create table with specified dimensions
        table = document.add_table(rows=rows, cols=cols)
        
        # Apply style if specified
        if style:
            try:
                table.style = style
            except KeyError:
                return f"Warning: Style '{style}' not found. Table added with default style."
        
        # Fill with data if provided
        if data:
            data_list = data.split(',')
            
            # Check if data matches table dimensions
            if len(data_list) > rows * cols:
                return f"Error: Number of data elements ({len(data_list)}) exceeds table dimensions ({rows}x{cols})."
                
            # Pad with empty strings if too few data elements
            if len(data_list) < rows * cols:
                data_list.extend([''] * (rows * cols - len(data_list)))
                
            # Fill table cells
            for i in range(rows):
                for j in range(cols):
                    cell_idx = i * cols + j
                    if cell_idx < len(data_list):
                        table.cell(i, j).text = data_list[cell_idx].strip()
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Table with {rows} rows and {cols} columns added successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error adding table: {str(e)}"

def merge_table_cells(doc_id: str, table_index: int, start_row: int, start_col: int, end_row: int, end_col: int) -> str:
    """Merges cells in a table.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        table_index (int): The index of the table in the document (0-based).
        start_row (int): The starting row index (0-based).
        start_col (int): The starting column index (0-based).
        end_row (int): The ending row index (0-based).
        end_col (int): The ending column index (0-based).
    """
    try:
        document = load_document(doc_id)
        
        # Check if document has tables
        if not document.tables or len(document.tables) <= table_index:
            return f"Error: Table index {table_index} is out of range. Document has {len(document.tables) if document.tables else 0} tables."
        
        table = document.tables[table_index]
        
        # Validate row and column indices
        if start_row < 0 or end_row >= len(table.rows) or start_col < 0:
            return "Error: Row or column index out of range."
        
        # Get first cell in the range
        first_cell = table.cell(start_row, start_col)
        
        # Get last cell in the range
        last_cell = table.cell(end_row, end_col)
        
        # Merge the cells
        first_cell.merge(last_cell)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Cells merged from ({start_row},{start_col}) to ({end_row},{end_col}) in table {table_index}."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error merging table cells: {str(e)}"

def get_table_data(doc_id: str, table_index: int, include_empty_cells: bool = True) -> str:
    """Retrieves a table's data as formatted text.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        table_index (int): The index of the table in the document (0-based).
        include_empty_cells (bool): Whether to include empty cells in the output.
    """
    try:
        document = load_document(doc_id)
        
        # Check if document has tables
        if not document.tables or len(document.tables) <= table_index:
            return f"Error: Table index {table_index} is out of range. Document has {len(document.tables) if document.tables else 0} tables."
        
        table = document.tables[table_index]
        
        # Get table data
        result = []
        for i, row in enumerate(table.rows):
            row_data = []
            
            # Handle cells before the first actual cell
            for _ in range(row.grid_cols_before):
                if include_empty_cells:
                    row_data.append("")
            
            # Handle actual cells
            for cell in row.cells:
                # Get all text from the cell, including from nested tables
                cell_text = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip() or include_empty_cells:
                        cell_text.append(paragraph.text)
                
                # Add nested tables as a note
                if cell.tables:
                    cell_text.append("[Contains nested table]")
                
                row_data.append("\n".join(cell_text))
            
            # Handle cells after the last actual cell
            for _ in range(row.grid_cols_after):
                if include_empty_cells:
                    row_data.append("")
            
            result.append(" | ".join(row_data))
        
        return "\n".join(result)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error retrieving table data: {str(e)}"

def list_tables(doc_id: str) -> str:
    """Lists all tables in a document with their dimensions.
    
    Args:
        doc_id (str): The document ID (filename without extension).
    """
    try:
        document = load_document(doc_id)
        
        if not document.tables:
            return f"No tables found in document '{doc_id}.docx'."
        
        tables_info = []
        for i, table in enumerate(document.tables):
            row_count = len(table.rows)
            # Get the maximum number of columns across all rows
            col_count = max([row.grid_cols_before + len(row.cells) + row.grid_cols_after 
                            for row in table.rows]) if row_count > 0 else 0
            
            first_cell_text = ""
            if row_count > 0 and len(table.rows[0].cells) > 0:
                first_cell_text = table.rows[0].cells[0].text[:30]
                if len(table.rows[0].cells[0].text) > 30:
                    first_cell_text += "..."
            
            # Get table style if available
            style_name = "Default"
            if hasattr(table, "style") and table.style:
                style_name = table.style.name
            
            tables_info.append(f"Table {i}: {row_count} rows x {col_count} columns. Style: '{style_name}'. First cell: '{first_cell_text}'")
        
        return "\n".join(tables_info)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error listing tables: {str(e)}"

def set_paragraph_properties(doc_id: str, paragraph_index: int, formatting: dict) -> str:
    """Sets various properties of a paragraph.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        paragraph_index (int): Index of the paragraph to format (0-based).
        formatting (dict): Dictionary with paragraph formatting options like:
            - alignment: 'LEFT', 'CENTER', 'RIGHT', 'JUSTIFY'
            - left_indent, right_indent, first_line_indent: Indentation in inches
            - space_before, space_after: Spacing in points
            - line_spacing: Line spacing as multiple or points
            - keep_together, keep_with_next, page_break_before, widow_control: Boolean pagination options
    """
    try:
        document = load_document(doc_id)
        
        if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
            return "Error: Paragraph index out of range."
        
        paragraph = document.paragraphs[paragraph_index]
        apply_paragraph_formatting(paragraph, formatting)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Paragraph {paragraph_index} properties set successfully."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error setting paragraph properties: {str(e)}"

def set_text_properties(doc_id: str, paragraph_index: int, run_index: int, formatting: dict) -> str:
    """Sets various properties of a run of text.
    
    Args:
        doc_id (str): The document ID (filename without extension).
        paragraph_index (int): Index of the paragraph containing the run (0-based).
        run_index (int): Index of the run within the paragraph (0-based).
        formatting (dict): Dictionary with text formatting options like:
            - name: Font name
            - size: Font size in points
            - bold, italic, underline: Boolean style options
            - color: Color as hex (#RRGGBB) or rgb(r,g,b)
    """
    try:
        document = load_document(doc_id)
        
        if paragraph_index < 0 or paragraph_index >= len(document.paragraphs):
            return "Error: Paragraph index out of range."
        
        paragraph = document.paragraphs[paragraph_index]
        
        if run_index < 0 or run_index >= len(paragraph.runs):
            return f"Error: Run index {run_index} is out of range. Paragraph has {len(paragraph.runs)} runs."
        
        run = paragraph.runs[run_index]
        apply_run_formatting(run, formatting)
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Text properties set for run {run_index} in paragraph {paragraph_index}."
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error setting text properties: {str(e)}"
