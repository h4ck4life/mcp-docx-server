"""
Document creation and content management operations.
"""

from docx import Document
import os
from mcp_docx_server.utils import get_document_path, load_document, add_content_to_document

def create_document(doc_id: str, title: str = "New Document") -> str:
    """Creates a new Word document with a title."""
    try:
        document = Document()
        document.add_heading(title, 0)
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        return f"Document '{doc_id}.docx' created successfully at path: {os.path.abspath(doc_path)}"
    except Exception as e:
        return f"Error creating document: {str(e)}"

def create_complete_document(doc_id: str, title: str = "New Document", content: list = None) -> str:
    """Creates a new Word document with title and content in a single operation."""
    try:
        document = Document()
        document.add_heading(title, 0)
        
        if not add_content_to_document(document, content):
            return "Error in table data: Number of data elements does not match table dimensions."
        
        doc_path = get_document_path(doc_id)
        document.save(doc_path)
        
        return f"Document '{doc_id}.docx' created successfully with title and {len(content) if content else 0} content items at path: {os.path.abspath(doc_path)}"
    except Exception as e:
        return f"Error creating document: {str(e)}"

def update_document(doc_id: str, title: str = None, content: list = None, append: bool = True) -> str:
    """Updates an existing Word document by appending or replacing content."""
    try:
        doc_path = get_document_path(doc_id)
        
        if not os.path.exists(doc_path) and append:
            return f"Document '{doc_id}.docx' does not exist and cannot be updated. Create it first."
        
        if not append or not os.path.exists(doc_path):
            document = Document()
            if title:
                document.add_heading(title, 0)
        else:
            document = Document(doc_path)
            if title:
                document.add_heading(title, 1)
        
        if not add_content_to_document(document, content):
            return "Error in table data: Number of data elements does not match table dimensions."
        
        document.save(doc_path)
        
        action = "updated by appending" if append else "replaced"
        title_msg = f" with new title" if title else ""
        content_msg = f" and {len(content) if content else 0} content items" if content else ""
        
        return f"Document '{doc_id}.docx' {action}{title_msg}{content_msg} successfully."
    except Exception as e:
        return f"Error updating document: {str(e)}"

def append_to_document(doc_id: str, content: list) -> str:
    """Appends content to an existing Word document."""
    return update_document(doc_id, title=None, content=content, append=True)

def replace_document(doc_id: str, title: str = None, content: list = None) -> str:
    """Replaces an existing Word document with new content."""
    return update_document(doc_id, title=title, content=content, append=False)

def read_document(doc_id: str) -> str:
    """Reads the entire content of a Word document."""
    try:
        document = load_document(doc_id)
        full_text = [paragraph.text for paragraph in document.paragraphs]
        return '\n'.join(full_text)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Unexpected error: {str(e)}"

def check_document_exists(doc_id: str) -> str:
    """Checks if a Word document exists and can be read."""
    doc_path = get_document_path(doc_id)
    try:
        if os.path.exists(doc_path):
            document = Document(doc_path)
            paragraph_count = len(document.paragraphs)
            return f"Document '{doc_id}.docx' exists and is readable at path: {os.path.abspath(doc_path)}. Contains {paragraph_count} paragraphs."
        else:
            return f"Document '{doc_id}.docx' does not exist at path: {os.path.abspath(doc_path)}"
    except Exception as e:
        return f"Document '{doc_id}.docx' exists but cannot be read: {str(e)}"

def list_available_documents() -> str:
    """Lists all Word documents (.docx files) available in the server directory."""
    try:
        script_dir = os.path.dirname(os.path.abspath(__file__))
        parent_dir = os.path.dirname(script_dir)  # Go up one level to the project root
        docx_files = [f.replace('.docx', '') for f in os.listdir(parent_dir) if f.endswith('.docx')]
        
        if not docx_files:
            return "No Word documents (.docx files) found in the server directory."
        
        doc_list = "\n".join([f"- {doc}" for doc in docx_files])
        return f"Available Word documents (without .docx extension):\n{doc_list}"
    except Exception as e:
        return f"Error listing documents: {str(e)}"

def convert_to_pdf(doc_id: str) -> str:
    """Converts a Word document to PDF format."""
    try:
        from docx2pdf import convert
        
        doc_path = get_document_path(doc_id)
        if not os.path.exists(doc_path):
            return f"Error: Document '{doc_id}.docx' not found."
        
        script_dir = os.path.dirname(os.path.abspath(__file__))
        parent_dir = os.path.dirname(script_dir)  # Go up one level to the project root
        pdf_path = os.path.join(parent_dir, f"{doc_id}.pdf")
        
        convert(doc_path, pdf_path)
        
        return f"Document successfully converted to PDF at: {os.path.abspath(pdf_path)}"
    except Exception as e:
        return f"Error converting document to PDF: {str(e)}"

def analyze_document_structure(doc_id: str) -> str:
    """Analyzes the structure of a document, showing paragraphs, runs, and tables.
    
    This tool provides a detailed breakdown of the document structure,
    which can be helpful for understanding how to target specific elements
    for modification.
    
    Args:
        doc_id (str): The document ID (filename without extension).
    """
    try:
        document = load_document(doc_id)
        
        structure = []
        structure.append(f"Document Structure Analysis for '{doc_id}.docx':")
        structure.append(f"Total paragraphs: {len(document.paragraphs)}")
        structure.append(f"Total tables: {len(document.tables)}")
        structure.append("\nParagraph Details:")
        
        for i, para in enumerate(document.paragraphs):
            if not para.text.strip():
                structure.append(f"  Paragraph {i}: [Empty paragraph]")
                continue
                
            style = para.style.name if para.style else "Default"
            run_count = len(para.runs)
            structure.append(f"  Paragraph {i}: Style='{style}', Runs={run_count}")
            structure.append(f"    Text: \"{para.text[:50]}{'...' if len(para.text) > 50 else ''}\"")
            
            if run_count > 0:
                structure.append(f"    Run details:")
                for j, run in enumerate(para.runs):
                    bold = "Bold" if run.bold else "Normal"
                    italic = "Italic" if run.italic else "Normal"
                    style_name = run.style.name if run.style else "Default"
                    structure.append(f"      Run {j}: Style='{style_name}', {bold}, {italic}, Text=\"{run.text[:30]}{'...' if len(run.text) > 30 else ''}\"")
        
        if document.tables:
            structure.append("\nTable Details:")
            for i, table in enumerate(document.tables):
                row_count = len(table.rows)
                col_count = max([row.grid_cols_before + len(row.cells) + row.grid_cols_after 
                                for row in table.rows]) if row_count > 0 else 0
                
                style_name = table.style.name if table.style else "Default"
                
                structure.append(f"  Table {i}: {row_count} rows x {col_count} columns")
                structure.append(f"    Style: {style_name}")
                
                # Show a preview of the first few cells
                if row_count > 0 and len(table.rows[0].cells) > 0:
                    structure.append(f"    Preview:")
                    max_preview_rows = min(3, row_count)
                    for r in range(max_preview_rows):
                        row = table.rows[r]
                        cell_texts = []
                        for cell in row.cells[:min(3, len(row.cells))]:
                            cell_text = cell.text[:20]
                            if len(cell.text) > 20:
                                cell_text += "..."
                            cell_texts.append(f"\"{cell_text}\"")
                        
                        additional = "..." if len(row.cells) > 3 else ""
                        structure.append(f"      Row {r}: {', '.join(cell_texts)}{additional}")
        
        return "\n".join(structure)
    except ValueError as e:
        return str(e)
    except Exception as e:
        return f"Error analyzing document structure: {str(e)}"
